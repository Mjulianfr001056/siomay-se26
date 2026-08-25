"""
Generator Lampiran SPK (PPL & PML) — Sensus Ekonomi 2026.

Port dari notebook 'generator/Generator Lampiran Kontrak Kerja PPL.ipynb'
dan 'Generator Lampiran Kontrak Kerja PML.ipynb', diadaptasi ke format input
NIK-keyed: input/00_input_lampiran_spk.xlsx

Sheet input:
  - data_mitra       : nik, nama_mitra
  - no_urut_spk      : nik, no_urut_spk
  - alokasi_supervisi: nik_ppl, nik_pml
  - alokasi_tugas    : idsubsls, nmsls, kdkec, kddesa, kdsls, kdsubsls,
                       nmdesa, nik_ppl, nik_pml
  - kode_kecamatan   : kode, kecamatan          (opsional — pemetaan nama)

Fungsi inti `iter_generate()` adalah *generator* yang melempar event
log/progres/file agar UI bisa menampilkan proses secara langsung,
dengan gaya log yang sama seperti di notebook.
"""
import math
import os
import re
from copy import deepcopy

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ── Skema wajib file input ─────────────────────────────────────────────
REQUIRED_SCHEMA = {
    "data_mitra": ["nik", "nama_mitra"],
    "no_urut_spk": ["nik", "no_urut_spk"],
    "alokasi_supervisi": ["nik_ppl", "nik_pml"],
    "alokasi_tugas": [
        "idsubsls", "nmsls", "kdkec", "kddesa", "kdsls", "kdsubsls",
        "nmdesa", "nik_ppl", "nik_pml",
    ],
}
OPTIONAL_SCHEMA = {
    "kode_kecamatan": ["kode", "kecamatan"],
}

# Nomor halaman pertama dokumen PML dalam buku (untuk penomoran header).
# Halaman 1 SPK PML = hal. 8 buku; halaman 2 (tabel wilayah) = hal. 9.
PAGE_START_PML = 9


def validate_input(file_path: str):
    """
    Validasi struktur file Excel input Lampiran SPK.
    Returns (is_valid: bool, errors: list[str], dfs: dict[str, DataFrame])
    """
    errors, dfs = [], {}
    try:
        xl = pd.ExcelFile(file_path)
    except Exception as e:
        return False, [f"Gagal membaca file Excel: {e}"], {}

    missing_sheets = [s for s in REQUIRED_SCHEMA if s not in xl.sheet_names]
    if missing_sheets:
        errors.append(
            f"Sheet wajib tidak ditemukan: {', '.join(missing_sheets)}"
        )
        return False, errors, {}

    for sheet, expected in {**REQUIRED_SCHEMA, **OPTIONAL_SCHEMA}.items():
        if sheet not in xl.sheet_names:
            continue
        df = xl.parse(sheet, dtype=str)
        df.columns = [str(c).strip() for c in df.columns]
        dfs[sheet] = df
        miss = [c for c in expected if c not in df.columns]
        if miss:
            errors.append(f"Sheet '{sheet}' kehilangan kolom: {', '.join(miss)}")

    if errors:
        return False, errors, dfs

    for sheet in REQUIRED_SCHEMA:
        if len(dfs[sheet]) == 0:
            errors.append(f"Sheet '{sheet}' kosong.")

    return (len(errors) == 0), errors, dfs


def _norm(v) -> str:
    """Normalisasi nilai sel menjadi string bersih."""
    s = "" if pd.isna(v) else str(v)
    return s.strip()


def load_context(dfs: dict) -> dict:
    """
    Susun konteks populasi dari DataFrame hasil validasi:
      mitra   : {nik: nama_mitra}
      urut    : {nik: no_urut_spk}
      tugas   : DataFrame alokasi_tugas (kolom ternormalisasi)
      kec_map : {kode: nama_kecamatan}  (kosong bila sheet tak ada)
    """
    mitra = {
        _norm(r["nik"]): _norm(r["nama_mitra"])
        for _, r in dfs["data_mitra"].iterrows()
        if _norm(r["nik"])
    }
    urut = {
        _norm(r["nik"]): _norm(r["no_urut_spk"])
        for _, r in dfs["no_urut_spk"].iterrows()
        if _norm(r["nik"])
    }
    tugas = dfs["alokasi_tugas"].copy()
    for col in tugas.columns:
        tugas[col] = tugas[col].map(_norm)
    kec_map = {}
    if "kode_kecamatan" in dfs:
        kc = dfs["kode_kecamatan"]
        kec_map = dict(zip(kc["kode"].map(_norm), kc["kecamatan"].map(_norm)))
    return {"mitra": mitra, "urut": urut, "tugas": tugas, "kec_map": kec_map}


# ── Fungsi manipulasi DOCX (port dari notebook) ────────────────────────


def replace_text_in_docx(doc: Document, replacements: dict) -> None:
    """
    Ganti placeholder {{key}} dengan value di semua paragraf dan tabel.
    Placeholder yang terpecah di beberapa run akan disatukan terlebih dahulu.

    Setelah penggantian, semua run berwarna merah pada template (misalnya
    placeholder {{jml_sls}} atau contoh nominal Rp) direset ke hitam agar
    dokumen output terlihat rapi.
    """
    def _replace_in_paragraphs(paragraphs):
        for para in paragraphs:
            full_text = "".join(run.text for run in para.runs)
            for placeholder, value in replacements.items():
                if placeholder in full_text:
                    full_text = full_text.replace(placeholder, str(value))
                    if para.runs:
                        para.runs[0].text = full_text
                        for run in para.runs[1:]:
                            run.text = ""

    _replace_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _replace_in_paragraphs(cell.paragraphs)

    # Reset semua run berwarna merah di seluruh dokumen (paragraf, tabel,
    # header & footer) menjadi hitam.
    _clean_red_colors(doc)


# ── Reset warna merah pada dokumen output ─────────────────────────────

def _is_red_hex(val: str) -> bool:
    """True bila *val* adalah heksadesimal RGB yang didominasi merah
    (mis. FF0000, EE0000, CC3333, dsb.)."""
    if not val or val.lower() in ("auto", "000000", "000"):
        return False
    try:
        r = int(val[0:2], 16)
        g = int(val[2:4], 16)
        b = int(val[4:6], 16)
    except (ValueError, IndexError):
        return False
    # Kanal merah tinggi; hijau & biru rendah
    return r > 128 and g < 100 and b < 100


def _clean_red_colors(doc: Document) -> None:
    """Hapus atribut warna merah dari semua <w:rPr> di dokumen.

    Template Lampiran SPK memakai warna FF0000/EE0000 untuk placeholder
    dan contoh data.  Setelah populasi, warna-warna tersebut direset ke
    hitam (default) agar dokumen output terlihat profesional.
    """
    def _clean_runs_in_paragraphs(paragraphs):
        for para in paragraphs:
            for run in para.runs:
                rpr = run._r.find(qn("w:rPr"))
                if rpr is None:
                    continue
                color_el = rpr.find(qn("w:color"))
                if color_el is None:
                    continue
                hex_val = color_el.get(qn("w:val"), "auto")
                if _is_red_hex(hex_val):
                    rpr.remove(color_el)

    _clean_runs_in_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _clean_runs_in_paragraphs(cell.paragraphs)
    for section in doc.sections:
        for hf in (section.header, section.footer,
                    section.first_page_header, section.first_page_footer,
                    section.even_page_header, section.even_page_footer):
            if hf is not None:
                _clean_runs_in_paragraphs(hf.paragraphs)


def _find_wilayah_table(doc: Document, need_nama_col: bool):
    """Cari tabel DAFTAR WILAYAH KERJA berdasarkan header 'KECAMATAN'."""
    table = None
    for t in doc.tables:
        if t.rows:
            header = " ".join(c.text for c in t.rows[0].cells)
            ok = "kecamatan" in header.lower()
            if need_nama_col:
                ok = ok and "nama" in header.lower()
            if ok:
                table = t
                break
    if table is None and need_nama_col and len(doc.tables) >= 2:
        table = doc.tables[1]  # fallback notebook PML: tabel kedua
    return table


def _detect_data_start(table):
    """Baris data pertama = setelah semua baris header (port notebook)."""
    data_start = 0
    markers = ("(1)", "(2)", "(3)", "(4)", "(5)", "No",
               "KECAMATAN/DISTRIK", "DESA/KAMPUNG/NAGARI",
               "DESA/ KAMPUNG/ NAGARI", "Nama Petugas Lapangan Sensus")
    for i, row in enumerate(table.rows):
        texts = [c.text.strip() for c in row.cells]
        if any(t in markers for t in texts):
            data_start = i + 1
    return data_start


def _make_cell(text, ref_cell, vmerge=None):
    """
    Buat elemen <w:tc> baru berdasarkan ref_cell.
    vmerge: None | 'restart' | 'continue'  (sesuai teknik notebook PML)
    """
    nc = deepcopy(ref_cell._tc)
    if vmerge is not None:
        tcPr = nc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            nc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:vMerge")):
            tcPr.remove(old)
        vm = OxmlElement("w:vMerge")
        if vmerge == "restart":
            vm.set(qn("w:val"), "restart")
        tcPr.append(vm)
    for p in nc.findall(qn("w:p")):
        for r in p.findall(qn("w:r")):
            p.remove(r)
        if vmerge == "continue":
            continue
        r_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        r_el.append(rpr)
        t_el = OxmlElement("w:t")
        t_el.text = str(text)
        t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        r_el.append(t_el)
        p.append(r_el)
    return nc


def _append_row(tbl_xml, ref_row, values, vmerges=None):
    new_tr = OxmlElement("w:tr")
    ref_trPr = ref_row._tr.find(qn("w:trPr"))
    if ref_trPr is not None:
        new_tr.append(deepcopy(ref_trPr))
    rc = ref_row.cells
    for idx, val in enumerate(values):
        vm = vmerges[idx] if vmerges else None
        new_tr.append(_make_cell(val, rc[idx], vmerge=vm))
    tbl_xml.append(new_tr)


def add_wilayah_rows(doc: Document, wilayah_rows: list, log=None) -> None:
    """
    Isi tabel DAFTAR WILAYAH KERJA dokumen PPL (4 kolom).
    wilayah_rows: list of (no, kecamatan_text, desa_text, jumlah_sls)
    """
    table = _find_wilayah_table(doc, need_nama_col=False)
    if table is None:
        if log:
            log("⚠️  Tabel DAFTAR WILAYAH KERJA tidak ditemukan!", "WARN")
        return
    tbl_xml = table._tbl
    all_rows = tbl_xml.findall(qn("w:tr"))
    data_start = _detect_data_start(table)
    for row_xml in all_rows[data_start:]:
        tbl_xml.remove(row_xml)
    ref_row = table.rows[-1]
    for (no, kec_text, desa_text, jml_sls) in wilayah_rows:
        _append_row(tbl_xml, ref_row, [no, kec_text, desa_text, jml_sls])


def add_pml_wilayah_rows(doc: Document, pml_rows: list, log=None) -> None:
    """
    Isi tabel DAFTAR PETUGAS DAN WILAYAH KERJA dokumen PML (5 kolom),
    dengan vMerge vertikal untuk kolom No & Nama PPL yang menempati
    beberapa baris (port persis dari notebook PML).
    pml_rows: list of (no_str, nama_ppl, kecamatan_text, desa_text, jml)
              nama_ppl hanya terisi pada baris pertama tiap kelompok PPL.
    """
    table = _find_wilayah_table(doc, need_nama_col=True)
    if table is None:
        if log:
            log("⚠️  Tabel DAFTAR PETUGAS DAN WILAYAH KERJA tidak ditemukan!",
                "WARN")
        return
    tbl_xml = table._tbl
    all_rows = tbl_xml.findall(qn("w:tr"))
    data_start = _detect_data_start(table)
    for row_xml in all_rows[data_start:]:
        tbl_xml.remove(row_xml)
    ref_row = table.rows[-1]

    # Hitung rowspan tiap kelompok PPL (baris berurutan bernama sama)
    span_map = []
    i = 0
    while i < len(pml_rows):
        no, nama, kec, desa, jml = pml_rows[i]
        span = 1
        j = i + 1
        while j < len(pml_rows) and pml_rows[j][1] == "":
            span += 1
            j += 1
        span_map.append((no, nama, kec, desa, jml, span))
        for k in range(1, span):
            no2, _, kec2, desa2, jml2 = pml_rows[i + k]
            span_map.append((no2, "", kec2, desa2, jml2, 0))  # continuation
        i += span

    for entry in span_map:
        no, nama, kec, desa, jml, span = entry
        is_first = span > 0
        is_continue = span == 0
        if is_first and span > 1:
            vm_no_nama = "restart"
        elif is_continue:
            vm_no_nama = "continue"
        else:
            vm_no_nama = None
        _append_row(
            tbl_xml, ref_row,
            [no, nama, kec, desa, jml],
            vmerges=[vm_no_nama, vm_no_nama, None, None, None],
        )


def _build_page_field_runs(prefix, suffix, start_page):
    """Kembalikan list elemen <w:r> yang membentuk prefix{PAGE}suffix."""
    runs = []
    r_pre = OxmlElement("w:r")
    t_pre = OxmlElement("w:t")
    t_pre.text = prefix
    r_pre.append(t_pre)
    runs.append(r_pre)

    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin.append(fc_begin)
    runs.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    r_instr.append(instr)
    runs.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    runs.append(r_sep)

    r_val = OxmlElement("w:r")
    t_val = OxmlElement("w:t")
    t_val.text = str(start_page)
    r_val.append(t_val)
    runs.append(r_val)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    runs.append(r_end)

    r_suf = OxmlElement("w:r")
    t_suf = OxmlElement("w:t")
    t_suf.text = suffix
    r_suf.append(t_suf)
    runs.append(r_suf)
    return runs


def fix_page_numbers(doc: Document, start_page: int, log=None) -> None:
    """
    FIX dari notebook PML: nomor halaman di HEADER berformat dinamis -PAGE-,
    rata tengah; multi-section lanjut menghitung (tidak restart).
    """
    pattern = re.compile(r"^(-+)\s*\d+\s*(-+)$")
    replaced = False

    for i, section in enumerate(doc.sections):
        pgNumType = section._sectPr.find(qn("w:pgNumType"))
        if i == 0:
            if pgNumType is None:
                pgNumType = OxmlElement("w:pgNumType")
                section._sectPr.append(pgNumType)
            pgNumType.set(qn("w:start"), str(start_page))
        elif pgNumType is not None and qn("w:start") in pgNumType.attrib:
            del pgNumType.attrib[qn("w:start")]

        hf_pairs = [
            (section.header, section.footer),
            (section.first_page_header, section.first_page_footer),
            (section.even_page_header, section.even_page_footer),
        ]
        for header, footer in hf_pairs:
            prefix, suffix = None, None
            if header is not None:
                for para in header.paragraphs:
                    m = pattern.match(
                        "".join(r.text for r in para.runs).strip()
                    )
                    if m:
                        prefix, suffix = m.group(1), m.group(2)
                        break
            if prefix is None and footer is not None:
                for para in footer.paragraphs:
                    m = pattern.match(
                        "".join(r.text for r in para.runs).strip()
                    )
                    if m:
                        prefix, suffix = m.group(1), m.group(2)
                        p_xml = footer._element
                        for p in p_xml.findall(qn("w:p")):
                            p_xml.remove(p)
                        break
            if prefix is None and suffix is None:
                prefix, suffix = "-", "-"
            if header is not None:
                header.is_linked_to_previous = False
                h_xml = header._element
                for p in h_xml.findall(qn("w:p")):
                    h_xml.remove(p)
                target_para = header.add_paragraph()
                target_para.alignment = 1  # CENTER
                for run_el in _build_page_field_runs(prefix, suffix,
                                                     start_page):
                    target_para._p.append(run_el)
                replaced = True

    if not replaced and log:
        log("⚠️  Penyesuaian nomor halaman selesai dengan mode default.",
            "WARN")


def _slug(name: str) -> str:
    return name.replace(" ", "_").replace("/", "-")


def _officer_list(ctx: dict, role_col: str):
    """Daftar petugas (urut sesuai sheet data_mitra) untuk satu peran."""
    tugas = ctx["tugas"]
    active = set(tugas[role_col]) - {""}
    seen, officers = set(), []
    for _, row in ctx["_mitra_df"].iterrows():
        nik = _norm(row["nik"])
        if not nik or nik in seen or nik not in active:
            continue
        seen.add(nik)
        nama = ctx["mitra"].get(nik) or nik
        no_urut = ctx["urut"].get(nik, "")
        officers.append((nik, nama, no_urut))
    return officers


def iter_generate(kind: str, ctx: dict, template_path: str, out_dir: str):
    """
    Generator populasi dokumen Lampiran SPK. kind: 'ppl' | 'pml'.
    Melempar event dict agar UI bisa streaming:
      {"t": "log", "level": str, "msg": str}
      {"t": "start", "total": int}
      {"t": "progress", "done": int, "total": int}
      {"t": "file", "path": str}
      {"t": "done", "generated": [str], "skipped": [str]}
    """
    os.makedirs(out_dir, exist_ok=True)
    tugas = ctx["tugas"]
    kec_map = ctx["kec_map"]
    mitra = ctx["mitra"]
    role_col = "nik_ppl" if kind == "ppl" else "nik_pml"
    emoji = "👤" if kind == "ppl" else "🔍"

    officers = _officer_list(ctx, role_col)
    yield {"t": "log", "level": "STEP",
           "msg": f"🚀 Memproses {len(officers)} "
                  f"{'petugas PPL…' if kind == 'ppl' else 'PML…'}"}

    generated, skipped = [], []
    total = len(officers)

    for idx, (nik, nama, no_urut) in enumerate(officers, start=1):
        yield {"t": "log", "level": "INFO",
               "msg": f"{emoji} [{no_urut or '-'}] {nama}"}

        df_x = tugas[tugas[role_col] == nik].copy()
        if df_x.empty:
            yield {"t": "log", "level": "WARN",
                   "msg": f"   ⚠️  Tidak ada data alokasi tugas untuk "
                          f"\"{nama}\" — dilewati"}
            skipped.append(nama)
            continue

        # ── Halaman 1: hitung SLS & target minimal 40% ────────────────
        jml_sls = (df_x["idsubsls"].nunique()
                   if "idsubsls" in df_x.columns else len(df_x))
        jml_sls_min = math.ceil(0.40 * jml_sls)

        stat = (f"   📊 Total SLS={jml_sls}  |  "
                f"Target 40% (min)={jml_sls_min}")
        if kind == "pml":
            jml_ppl = df_x["nik_ppl"].nunique()
            stat += f"  |  PPL={jml_ppl}"
        yield {"t": "log", "level": "INFO", "msg": stat}

        def kec_text(kode: str) -> str:
            nm = kec_map.get(kode, "")
            return f"[{kode}] {nm}" if nm else f"[{kode}]"

        # ── Halaman 2: susun baris tabel wilayah kerja ────────────────
        if kind == "ppl":
            grp_cols = ["kdkec", "kddesa", "nmdesa"]
            grp = (
                df_x.groupby(grp_cols, sort=True)["idsubsls"].nunique()
                .reset_index(name="jml_subsls")
                if "idsubsls" in df_x.columns
                else df_x.groupby(grp_cols, sort=True).size()
                         .reset_index(name="jml_subsls")
            )
            rows_wilayah = [
                (f"{i}.",
                 kec_text(r.kdkec),
                 f"[{r.kddesa}] {r.nmdesa}",
                 str(r.jml_subsls))
                for i, r in enumerate(grp.itertuples(), start=1)
            ]
            yield {"t": "log", "level": "INFO",
                   "msg": f"   🗺️  {len(rows_wilayah)} desa/kelurahan "
                          f"dalam wilayah kerja"}
        else:
            grp_cols = ["nik_ppl", "kdkec", "kddesa", "nmdesa"]
            grp = (
                df_x.groupby(grp_cols, sort=True)["idsubsls"].nunique()
                .reset_index(name="jml_subsls")
                if "idsubsls" in df_x.columns
                else df_x.groupby(grp_cols, sort=True).size()
                         .reset_index(name="jml_subsls")
            )
            rows_wilayah = []
            row_no, prev_ppl = 1, None
            for r in grp.itertuples():
                nik_ppl = r.nik_ppl
                nama_ppl = mitra.get(nik_ppl, nik_ppl)
                display = nama_ppl if nama_ppl != prev_ppl else ""
                prev_ppl = nama_ppl
                rows_wilayah.append((
                    f"{row_no}." if display else "",
                    display,
                    kec_text(r.kdkec),
                    f"[{r.kddesa}] {r.nmdesa}",
                    str(r.jml_subsls),
                ))
                if display:
                    row_no += 1
            yield {"t": "log", "level": "INFO",
                   "msg": f"   🗺️  {len(rows_wilayah)} baris wilayah "
                          f"({jml_ppl} PPL)"}

# ── Isi template ──────────────────────────────────────────────
        doc = Document(template_path)
        replace_text_in_docx(doc, {
            "{{nomor_urut_spk}}": no_urut,
            "{{jml_sls}}": str(jml_sls),
            "{{jml_sls_min}}": str(jml_sls_min),
        })
        if kind == "ppl":
            add_wilayah_rows(doc, rows_wilayah,
                             log=lambda m, lv="WARN": None)
        else:
            add_pml_wilayah_rows(doc, rows_wilayah,
                                 log=lambda m, lv="WARN": None)
            fix_page_numbers(doc, PAGE_START_PML)

        safe_nama = _slug(nama)
        fname = f"{no_urut or idx}_{safe_nama}.docx"
        out_path = os.path.join(out_dir, fname)
        doc.save(out_path)
        generated.append(out_path)
        yield {"t": "file", "path": out_path}
        yield {"t": "log", "level": "OK",
               "msg": f"   ✅ Dokumen dibuat: {fname}"}
        yield {"t": "progress", "done": idx, "total": total}

    yield {"t": "log", "level": "STEP",
           "msg": "─" * 50}
    yield {"t": "log",
           "level": "OK" if generated else "ERROR",
           "msg": f"Selesai: {len(generated)} dokumen berhasil, "
                  f"{len(skipped)} dilewati."}
    if skipped:
        yield {"t": "log", "level": "WARN",
               "msg": "Dilewati: " + ", ".join(skipped)}
    yield {"t": "done", "generated": generated, "skipped": skipped}


def prepare_context(dfs: dict) -> dict:
    """load_context + simpan urutan asli sheet data_mitra."""
    ctx = load_context(dfs)
    ctx["_mitra_df"] = dfs["data_mitra"]
    ctx["_sheets"] = set(dfs.keys())
    return ctx





