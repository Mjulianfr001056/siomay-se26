"""
Generator BAPP Termin 2 (PPL) - Sensus Ekonomi 2026.

Port dari notebook 'generator/Generator_BAPP_PPL_grid_bukti_dukung (1).ipynb'.

Mengisi template BAPP (Berita Acara Pemeriksaan Hasil Pekerjaan) Termin II
untuk Petugas Lapangan (PPL) berdasarkan data Excel, menyisipkan screenshot
bukti dukung sebagai grid adaptif atau pada halaman khusus.

Input Excel:
  Sheet 'input': nik, nama_lengkap, no_spk, no_urut_bapp_t2,
                 jml_sls_t2, bukti_dukung_bapp_t2

Template DOCX placeholders:
  {{no_urut_bapp_t2}}, {{no_spk}}, {{nama_lengkap}}, {{nik}},
  {{jml_sls_t2}}, {{bukti_dukung_bapp_t2}}
"""
import os
import re

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from utils.images import (
    HAS_HEIF,
    HAS_PIL,
    download_drive_evidence as _download_drive_evidence,
)
from utils.evidence import (
    IMAGE_ORIENTATION_AUTOMATIC,
    IMAGE_ORIENTATION_LANDSCAPE,
    IMAGE_ORIENTATION_PORTRAIT,
    insert_evidence as _insert_evidence,
)


# -- Skema input Excel ------------------------------------------------
SHEET_NAME = "input"

REQUIRED_COLUMNS = [
    "nik",
    "nama_lengkap",
    "no_spk",
    "no_urut_bapp_t2",
    "jml_sls_t2",
    "bukti_dukung_bapp_t2",
]

# Mapping: template placeholder (without braces) -> input column name
# Template uses: {{no_urut_bapp_t2}}, {{no_spk}}, {{nama_lengkap}},
#                {{nik}}, {{jml_sls_t2}}, {{bukti_dukung_bapp_t2}}
PLACEHOLDER_MAP = {
    "no_urut_bapp_t2": "no_urut_bapp_t2",
    "no_spk":          "no_spk",
    "nama_lengkap":    "nama_lengkap",
    "nik":             "nik",
    "jml_sls_t2":      "jml_sls_t2",
}

BUKTI_PLACEHOLDER = "{{bukti_dukung_bapp_t2}}"
LINK_COLUMN = "bukti_dukung_bapp_t2"

IMAGE_LAYOUT_GRID = "grid"
IMAGE_LAYOUT_DEDICATED_PAGES = "dedicated_pages"
DEDICATED_MAX_WIDTH_IN = 7.5
DEDICATED_MAX_HEIGHT_IN = 4.0


def validate_input(file_path: str):
    """
    Validasi struktur file Excel input BAPP T2 PPL.
    Returns (is_valid, errors, dfs) -- dfs = {sheet_name: DataFrame}.
    """
    errors, dfs = [], {}
    try:
        xl = pd.ExcelFile(file_path)
    except Exception as e:
        return False, [f"Gagal membaca file Excel: {e}"], {}

    if SHEET_NAME not in xl.sheet_names:
        errors.append(
            f"Sheet '{SHEET_NAME}' tidak ditemukan. "
            f"Sheet yang tersedia: {', '.join(xl.sheet_names)}"
        )
        return False, errors, {}

    df = xl.parse(SHEET_NAME, dtype=str)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.fillna("")
    dfs[SHEET_NAME] = df

    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        errors.append(
            f"Sheet '{SHEET_NAME}' kehilangan kolom: {', '.join(missing)}"
        )
        return False, errors, dfs

    if len(df) == 0:
        errors.append(f"Sheet '{SHEET_NAME}' kosong.")
        return False, errors, dfs

    return True, [], dfs


def _norm(v) -> str:
    """Normalisasi nilai sel menjadi string bersih."""
    if pd.isna(v):
        return ""
    return str(v).strip()


def _format_no_urut_bapp_t2(v) -> str:
    """Pertahankan nomor urut BAPP T2 sebagai tiga digit (mis. ``001``).

    Excel dapat menyimpan masukan/format tampilan ``001`` sebagai angka ``1``.
    Setelah dibaca pandas, informasi nol di depannya tidak lagi tersedia, jadi
    nomor urut numerik dinormalisasi kembali menjadi sedikitnya tiga digit.
    """
    value = _norm(v)
    return value.zfill(3) if value.isdigit() else value


def generate_input_template(file_path: str):
    """Buat template Excel input BAPP T1 PML (sample data)."""
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = SHEET_NAME

    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill(
        start_color="4F81BD", end_color="4F81BD", fill_type="solid"
    )
    centered = Alignment(horizontal="center", vertical="center")
    thin_border = Border(
        left=Side(style="thin", color="BFBFBF"),
        right=Side(style="thin", color="BFBFBF"),
        top=Side(style="thin", color="BFBFBF"),
        bottom=Side(style="thin", color="BFBFBF"),
    )
    left_align = Alignment(horizontal="left", vertical="center")

    for col_idx, col_name in enumerate(REQUIRED_COLUMNS, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = centered
        cell.border = thin_border

    sample = {
        "nik": "6304xxxx",
        "nama_lengkap": "Nama Petugas",
        "no_spk": "B-001/SPK-PML-SE2026/6304/PL.200/2026",
        "no_urut_bapp_t2": "001",
        "jml_sls_t2": "12",
        "bukti_dukung_bapp_t2": "https://drive.google.com/open?id=XXXXX",
    }
    for col_idx, col_name in enumerate(REQUIRED_COLUMNS, 1):
        cell = ws.cell(row=2, column=col_idx, value=sample.get(col_name, ""))
        cell.alignment = left_align
        cell.border = thin_border

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        letter = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[letter].width = max(max_len + 3, 12)

    wb.save(file_path)


def replace_text_preserving_runs(doc: Document, replacements: dict) -> None:
    """
    Ganti placeholder {{key}} sambil mempertahankan formatting tiap run.

    Pendekatan:
    1. Bangun peta karakter: char_idx -> (run_idx, offset).
    2. Temukan span placeholder di full_text.
    3. Per placeholder (dari kanan ke kiri agar indeks tidak bergeser):
       - Tulis nilai ke run pertama yang menyentuh span.
       - Bersihkan teks dari run-run lain dalam span.
       - Hapus warna editorial (biru/merah) dan underline dari run placeholder.
    """

    def _strip_placeholder_fmt(run):
        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            return
        color = rpr.find(qn("w:color"))
        if color is not None:
            val = color.get(qn("w:val")) or ""
            if val.lower() in ("0000ff", "ff0000"):
                rpr.remove(color)
        u = rpr.find(qn("w:u"))
        if u is not None:
            rpr.remove(u)

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            runs = para.runs
            if not runs:
                continue
            while True:
                full_text = "".join(run.text for run in runs)
                match = next(
                    (
                        m for m in re.finditer(r"\{\{[^}]+\}\}", full_text)
                        if m.group(0) in replacements
                    ),
                    None,
                )
                if match is None:
                    break

                run_starts = []
                position = 0
                for run in runs:
                    run_starts.append(position)
                    position += len(run.text)
                first_ri = max(
                    i for i, start in enumerate(run_starts) if start <= match.start()
                )
                last_ri = max(
                    i for i, start in enumerate(run_starts) if start < match.end()
                )
                prefix = runs[first_ri].text[
                    :match.start() - run_starts[first_ri]
                ]
                suffix = runs[last_ri].text[
                    match.end() - run_starts[last_ri]:
                ]
                runs[first_ri].text = (
                    prefix + str(replacements[match.group(0)]) + suffix
                )
                _strip_placeholder_fmt(runs[first_ri])
                for ri in range(first_ri + 1, last_ri + 1):
                    runs[ri].text = ""

    def _process_table(table):
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)
                for nested_table in cell.tables:
                    _process_table(nested_table)

    _process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        _process_table(table)
    for section in doc.sections:
        _process_paragraphs(section.header.paragraphs)
        for table in section.header.tables:
            _process_table(table)
        _process_paragraphs(section.footer.paragraphs)
        for table in section.footer.tables:
            _process_table(table)


def _extract_file_id(link: str):
    """Ambil file ID dari tautan Google Drive."""
    link = link.strip()
    if not link:
        return None
    m = re.search(r"[?&]id=([-\w]+)", link)
    if m:
        return m.group(1)
    m = re.search(r"/d/([-\w]+)", link)
    if m:
        return m.group(1)
    m = re.search(r"[-\w]{25,}", link)
    if m:
        return m.group(0)
    return None


def insert_gdrive_images(doc: Document, links_str: str,
                         placeholder: str = None,
                         image_layout: str = IMAGE_LAYOUT_GRID,
                         image_orientation: str = IMAGE_ORIENTATION_PORTRAIT):
    """Sisipkan gambar/PDF; tiap halaman PDF selalu memakai halaman khusus."""
    return _insert_evidence(
        doc,
        links_str,
        placeholder or BUKTI_PLACEHOLDER,
        image_layout,
        _extract_file_id,
        replace_text_preserving_runs,
        _download_drive_evidence,
        image_orientation,
    )


def _slug(name: str) -> str:
    s = re.sub(r"[^A-Za-z0-9]+", "_", str(name).strip())
    return s.strip("_")[:40] or "tanpa_nama"


def iter_generate(dfs: dict, template_path: str, out_dir: str,
                  image_layout: str = IMAGE_LAYOUT_GRID,
                  image_orientation: str = IMAGE_ORIENTATION_PORTRAIT):
    """
    Generator populasi dokumen BAPP T2 PPL.

    Yields event dicts:
      {"t": "log",      "level": str, "msg": str}
      {"t": "progress",  "done": int, "total": int}
      {"t": "file",      "path": str}
      {"t": "done",      "generated": [str], "skipped": [str]}
    """
    os.makedirs(out_dir, exist_ok=True)
    df = dfs.get(SHEET_NAME)
    if df is None or df.empty:
        yield {"t": "log", "level": "ERROR",
               "msg": f"Sheet '{SHEET_NAME}' kosong atau tidak ada."}
        yield {"t": "done", "generated": [], "skipped": []}
        return

    total = len(df)
    generated, skipped = [], []

    yield {"t": "log", "level": "STEP",
           "msg": f"Memproses {total} PPL..."}
    if not HAS_PIL:
        yield {"t": "log", "level": "WARN",
               "msg": "Pillow tidak terinstal - "
                      "sisipkan screenshot dilewati."}
    if HAS_HEIF:
        yield {"t": "log", "level": "INFO",
               "msg": "HEIC/HEIF (foto iPhone) didukung."}

    for idx, row in df.iterrows():
        nik = _norm(row.get("nik", ""))
        nama = _norm(row.get("nama_lengkap", ""))
        link_gd = _norm(row.get(LINK_COLUMN, ""))

        if not nik:
            yield {"t": "log", "level": "WARN",
                   "msg": f"   Baris {idx + 1}: NIK kosong - dilewati"}
            skipped.append(f"baris {idx + 1} (NIK kosong)")
            continue

        who = nama or nik
        yield {"t": "log", "level": "INFO",
               "msg": f"[{idx + 1}/{total}] {who} (NIK {nik})"}

        doc = Document(template_path)

        # Build replacements from PLACEHOLDER_MAP
        replacements = {}
        for ph_key, input_col in PLACEHOLDER_MAP.items():
            val = _norm(row.get(input_col, ""))
            if input_col == "no_urut_bapp_t2":
                val = _format_no_urut_bapp_t2(val)
            replacements["{{" + ph_key + "}}"] = val

        replace_text_preserving_runs(doc, replacements)

        n_img, img_warnings = insert_gdrive_images(
            doc, link_gd, image_layout=image_layout,
            image_orientation=image_orientation,
        )
        if n_img:
            yield {"t": "log", "level": "INFO",
                   "msg": f"   {n_img} screenshot disisipkan"}
        for w in img_warnings:
            yield {"t": "log", "level": "WARN", "msg": f"   {w}"}

        # Simpan DOCX
        safe_name = _slug(who)
        out_name = f"BAPP_PPL_Termin2_{idx + 1:03d}_{nik}_{safe_name}.docx"
        out_path = os.path.join(out_dir, out_name)
        doc.save(out_path)
        generated.append(out_path)
        yield {"t": "file", "path": out_path}
        yield {"t": "log", "level": "OK",
               "msg": f"   Tersimpan: {out_name}"}
        yield {"t": "progress", "done": idx + 1, "total": total}

    yield {"t": "log", "level": "STEP", "msg": "=" * 50}
    yield {"t": "log",
           "level": "OK" if generated else "ERROR",
           "msg": f"Selesai: {len(generated)} dokumen berhasil, "
                  f"{len(skipped)} dilewati."}
    if skipped:
        yield {"t": "log", "level": "WARN",
               "msg": "Dilewati: " + ", ".join(skipped)}
    yield {"t": "done", "generated": generated, "skipped": skipped}