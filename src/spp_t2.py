"""Generator Surat Pernyataan Penyelesaian (SPP) termin 2.

The supplied termin 2 workbook has one data sheet.  Its identity fields are
validated explicitly; any additional field used by the selected Word template
must be provided as a column bearing the same name.
"""
import os
import re

import pandas as pd
from docx import Document

from src.spp import replace_text_preserving_runs


SHEET_NAME = "input"
REQUIRED_COLUMNS = ["nik", "nama_lengkap", "no_spk", "no_input_spp_t2"]


def _norm(value) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def _template_fields(template_path: str) -> set[str]:
    """Return placeholder names, including fields located inside tables."""
    doc = Document(template_path)
    text = []

    def visit_table(table):
        for row in table.rows:
            for cell in row.cells:
                text.extend(paragraph.text for paragraph in cell.paragraphs)
                for nested in cell.tables:
                    visit_table(nested)

    text.extend(paragraph.text for paragraph in doc.paragraphs)
    for table in doc.tables:
        visit_table(table)
    for section in doc.sections:
        text.extend(paragraph.text for paragraph in section.header.paragraphs)
        text.extend(paragraph.text for paragraph in section.footer.paragraphs)
    return set(re.findall(r"\{\{\s*(\w+)\s*\}\}", "\n".join(text)))


def validate_input(file_path: str, template_path: str | None = None):
    """Validate the termin 2 SPP workbook and optional template field mapping."""
    try:
        with pd.ExcelFile(file_path) as workbook:
            if SHEET_NAME not in workbook.sheet_names:
                return False, [
                    f"Sheet '{SHEET_NAME}' tidak ditemukan. Sheet yang tersedia: "
                    f"{', '.join(workbook.sheet_names)}"
                ], {}
            df = workbook.parse(SHEET_NAME, dtype=str)
    except Exception as exc:
        return False, [f"Gagal membaca file Excel: {exc}"], {}
    df.columns = [str(column).strip() for column in df.columns]
    df = df.fillna("")
    errors = []
    missing = [column for column in REQUIRED_COLUMNS if column not in df.columns]
    if missing:
        errors.append("Sheet 'input' kekurangan kolom: " + ", ".join(missing))
    if df.empty:
        errors.append("Sheet 'input' kosong.")

    if template_path and not errors:
        fields = _template_fields(template_path)
        unmapped = sorted(fields - set(df.columns))
        if unmapped:
            errors.append(
                "Kolom input untuk placeholder template tidak ditemukan: "
                + ", ".join(unmapped)
            )
    return not errors, errors, {SHEET_NAME: df}


def iter_generate(kind: str, dfs: dict, template_path: str, out_dir: str):
    """Generate one SPP termin 2 DOCX per row for the selected officer role."""
    df = dfs.get(SHEET_NAME)
    if df is None or df.empty:
        yield {"t": "log", "level": "ERROR", "msg": "Sheet 'input' kosong atau tidak ada."}
        yield {"t": "done", "generated": []}
        return

    os.makedirs(out_dir, exist_ok=True)
    generated = []
    role = kind.upper()
    total = len(df)
    yield {"t": "log", "level": "STEP", "msg": f"Memproses {total} SPP termin 2 {role}..."}

    for index, row in df.iterrows():
        nik = _norm(row.get("nik", ""))
        nama = _norm(row.get("nama_lengkap", ""))
        if not nik:
            yield {"t": "log", "level": "WARN", "msg": f"Baris {index + 1}: NIK kosong - dilewati."}
            yield {"t": "progress", "done": index + 1, "total": total}
            continue

        doc = Document(template_path)
        replacements = {
            "{{" + column + "}}": _norm(value)
            for column, value in row.items()
        }
        replace_text_preserving_runs(doc, replacements)
        safe_name = re.sub(r"[^A-Za-z0-9]+", "_", nama or nik).strip("_")[:40] or "tanpa_nama"
        output_name = f"SPP_{role}_Termin2_{index + 1:03d}_{nik}_{safe_name}.docx"
        output_path = os.path.join(out_dir, output_name)
        doc.save(output_path)
        generated.append(output_path)
        yield {"t": "file", "path": output_path}
        yield {"t": "log", "level": "OK", "msg": f"Tersimpan: {output_name}"}
        yield {"t": "progress", "done": index + 1, "total": total}

    yield {"t": "done", "generated": generated}