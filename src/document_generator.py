"""
DOCX template filling engine.

Populates a .docx template for one data row by replacing ``{{field_name}}``
placeholders found in body paragraphs, table cells, headers and footers.

If python-docx is unavailable, callers can fall back to copy_template().
"""
import os
import re
import shutil
from copy import copy

try:
    from docx import Document
    HAS_DOCX = True
except Exception:  # pragma: no cover - depends on environment
    Document = None
    HAS_DOCX = False

# Placeholder names are deliberately ASCII and exact. Malformed markers (for
# example ``{{full name}}`` or ``{{ name }}``) are ordinary document text.
PLACEHOLDER_RE = re.compile(r"\{\{([A-Za-z0-9_]+)\}\}")

_EMPTY_TOKENS = {"", "nan", "none", "<na>", "nat"}


def clean_value(v) -> str:
    """Normalize a cell value into a clean display string."""
    if v is None:
        return ""
    try:
        if isinstance(v, float) and v != v:  # NaN
            return ""
    except Exception:
        pass
    s = str(v).strip()
    return "" if s.lower() in _EMPTY_TOKENS else s


def build_values(record: dict) -> dict:
    """Convert a dataframe record into {key: clean_str}, dropping empties."""
    return {
        str(k): clean_value(v)
        for k, v in record.items()
        if clean_value(v) != ""
    }


def iter_headers_and_footers(doc):
    """Yield every distinct primary, first-page, and even-page header/footer."""
    seen = set()
    story_names = (
        "header", "footer",
        "first_page_header", "first_page_footer",
        "even_page_header", "even_page_footer",
    )
    for section in doc.sections:
        for story_name in story_names:
            story = getattr(section, story_name)
            # Multiple sections commonly link to the same header/footer part.
            # Process each XML story once while still following those links.
            story_id = id(story._element)
            if story_id in seen:
                continue
            seen.add(story_id)
            yield story


def _iter_paragraphs(doc):
    """Yield paragraphs in the body and every Word header/footer variant."""
    def iter_table_paragraphs(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from cell.paragraphs
                    yield from iter_table_paragraphs(cell.tables)

    for p in doc.paragraphs:
        yield p
    yield from iter_table_paragraphs(doc.tables)
    for story in iter_headers_and_footers(doc):
        for p in story.paragraphs:
            yield p
        yield from iter_table_paragraphs(story.tables)


def extract_template_placeholders(template_path: str) -> set[str]:
    """Return all unique ``{{field_name}}`` keys used by a DOCX template.

    This uses the same paragraph traversal as :func:`fill_row`, so validation
    includes placeholders in document tables and all primary, first-page, and
    even-page headers/footers as well as the body. It also works when a
    placeholder is split across Word runs.
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx tidak tersedia")

    doc = Document(template_path)
    placeholders = set()
    for paragraph in _iter_paragraphs(doc):
        placeholders.update(PLACEHOLDER_RE.findall(paragraph.text))
    return placeholders


def validate_template_placeholders(
    template_path: str, expected_placeholders: set[str],
) -> dict:
    """Compare a template's placeholders with the expected placeholder set.

    Returns ``is_valid``, plus sorted ``missing``, ``unexpected`` and ``custom``
    keys. Built-in fields are mandatory; additional well-formed fields are
    accepted as custom Excel-backed placeholders.
    """
    actual = extract_template_placeholders(template_path)
    expected = set(expected_placeholders)
    return {
        "is_valid": expected <= actual,
        "missing": sorted(expected - actual),
        "unexpected": sorted(actual - expected),
        "custom": sorted(actual - expected),
    }


def custom_template_placeholders(
    template_path: str | None, builtin_template_path: str | None,
) -> list[str]:
    """Return custom keys added to *template_path*, in deterministic order."""
    if not template_path or not builtin_template_path:
        return []
    return sorted(
        extract_template_placeholders(template_path)
        - extract_template_placeholders(builtin_template_path)
    )


def validate_custom_columns(dfs: dict, sheet_name: str, custom_fields) -> list[str]:
    """Return validation errors for custom fields absent from one input sheet."""
    fields = list(custom_fields or [])
    if not fields:
        return []
    dataframe = dfs.get(sheet_name)
    if dataframe is None:
        return [f"Sheet '{sheet_name}' tidak tersedia untuk kolom kustom."]
    headers = {str(column).strip() for column in dataframe.columns}
    missing = [field for field in fields if field not in headers]
    if not missing:
        return []
    return [
        f"Sheet '{sheet_name}' kekurangan kolom placeholder kustom: "
        + ", ".join(missing)
    ]


def row_placeholder_replacements(row, fields=None) -> dict:
    """Build exact token replacements from a dataframe row, including blanks."""
    keys = fields if fields is not None else row.index
    return {
        "{{" + str(key) + "}}": clean_value(row.get(key, ""))
        for key in keys
        if re.fullmatch(r"[A-Za-z0-9_]+", str(key))
    }


def extend_input_template(
    source_path: str, output_path: str, sheet_name: str, custom_fields,
) -> str:
    """Copy an XLSX and append custom Text-format columns to *sheet_name*."""
    import openpyxl

    workbook = openpyxl.load_workbook(source_path)
    try:
        if sheet_name not in workbook.sheetnames:
            raise ValueError(f"Sheet '{sheet_name}' tidak ditemukan.")
        worksheet = workbook[sheet_name]
        headers = {
            str(cell.value).strip(): cell.column
            for cell in worksheet[1]
            if cell.value is not None
        }
        for field in custom_fields or []:
            if field in headers:
                continue
            column = worksheet.max_column + 1
            cell = worksheet.cell(row=1, column=column, value=field)
            if column > 1:
                source = worksheet.cell(row=1, column=column - 1)
                cell.font = copy(source.font)
                cell.fill = copy(source.fill)
                cell.border = copy(source.border)
                cell.alignment = copy(source.alignment)
                cell.protection = copy(source.protection)
                worksheet.column_dimensions[cell.column_letter].width = max(
                    worksheet.column_dimensions[source.column_letter].width or 0,
                    len(field) + 2,
                )
            for row_number in range(2, worksheet.max_row + 1):
                worksheet.cell(row=row_number, column=column).number_format = "@"
            cell.number_format = "@"
            headers[field] = column
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        workbook.save(output_path)
    finally:
        workbook.close()
    return output_path


def _replace_in_paragraph(para, values: dict):
    """
    Replace placeholders in one paragraph. Returns list of keys actually
    filled (non-empty value available).
    """
    text = para.text
    if "{{" not in text:
        return []

    filled = []

    def repl(match):
        key = match.group(1)
        val = values.get(key)
        if val:  # non-empty string
            filled.append(key)
            return val
        return match.group(0)  # keep token so it is reported unresolved

    new_text = PLACEHOLDER_RE.sub(repl, text)
    if new_text != text:
        runs = para.runs
        if runs:
            runs[0].text = new_text          # keep first-run formatting
            for r in runs[1:]:
                r.text = ""
        else:
            para.text = new_text
    return filled


def fill_row(template_path: str, values: dict, out_path: str) -> dict:
    """
    Fill one document. Returns:
      {"filled": [keys substituted], "unresolved": [tokens left in doc]}
    """
    if not HAS_DOCX:
        raise RuntimeError("python-docx tidak tersedia")

    doc = Document(template_path)
    paragraphs = list(_iter_paragraphs(doc))

    filled = set()
    for p in paragraphs:
        filled.update(_replace_in_paragraph(p, values))

    remaining = set()
    for p in paragraphs:
        remaining.update(PLACEHOLDER_RE.findall(p.text))

    unresolved = sorted(remaining - filled)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)
    return {"filled": sorted(filled), "unresolved": unresolved}


def copy_template(template_path: str, out_path: str) -> str:
    """Fallback: copy template as-is (used when python-docx is missing)."""
    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    shutil.copyfile(template_path, out_path)
    return out_path


def slugify(name: str, fallback: str = "tanpa_nama") -> str:
    """Filesystem-safe fragment from a person/unit name."""
    s = "".join(ch if ch.isalnum() else "_" for ch in str(name).strip())
    s = re.sub(r"_+", "_", s).strip("_")
    return s[:40] if s else fallback
