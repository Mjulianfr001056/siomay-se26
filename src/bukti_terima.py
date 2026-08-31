"""
Generator Bukti Terima Paket Internet — SE2026.

Port dari notebook 'generator/Generator_Bukti_Terima_Juli (7).ipynb'.

Membuat dokumen A4 berisi grid 2x2 tetap per halaman dari dokumen kosong
(bukan dari template .docx). Tiap sel berisi:
  - Nama lengkap (bold, centered)
  - Foto dari Google Drive (lazy-download via requests)
  - Detail: NIK, Wilayah Tugas, Operator, No Telp

Input Excel:
  Sheet 'clean': timestamp, nik, no_telp, nama_lengkap,
                 wilayah_tugas, operator, link_bukti_terima

Output: satu DOCX (banyak halaman, grid 2x2 tetap, tanpa border).
"""
import datetime
import io
import os
import re
import time

import pandas as pd
import requests

try:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from PIL import Image, ImageFile
    ImageFile.LOAD_TRUNCATED_IMAGES = True
    HAS_PIL = True
except ImportError:
    HAS_PIL = False

try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
    HAS_HEIF = True
except ImportError:
    HAS_HEIF = False


# ── Skema input Excel ──────────────────────────────────────────────────────
SHEET_NAME = "clean"

REQUIRED_COLUMNS = [
    "nik",
    "nama_lengkap",
    "wilayah_tugas",
    "operator",
    "no_telp",
    "link_bukti_terima",
]

# ── Layout A4 (cm) ─────────────────────────────────────────────────────────
MARGIN_CM           = 1.5
PAGE_W_CM           = 21.0
PAGE_H_CM           = 29.7
PARAGRAPH_BUFFER_CM = 0.5

USABLE_W_CM  = PAGE_W_CM - 2 * MARGIN_CM            # 18.0
USABLE_H_CM  = PAGE_H_CM - 2 * MARGIN_CM - PARAGRAPH_BUFFER_CM  # 26.2

COL_WIDTH_CM  = USABLE_W_CM / 2                      # 9.0
ROW_HEIGHT_CM = USABLE_H_CM / 2                      # 13.1

CELL_PADDING_CM   = 0.5
TEXT_ALLOWANCE_CM = 2.6

IMAGE_BOX_WIDTH_CM  = COL_WIDTH_CM - CELL_PADDING_CM            # 8.5
IMAGE_BOX_HEIGHT_CM = max(ROW_HEIGHT_CM - TEXT_ALLOWANCE_CM, 3.0)

_MAX_RETRIES   = 3
_RETRY_DELAY_S = 2


# ── Utilitas Google Drive ──────────────────────────────────────────────────

def _extract_file_id(link: str):
    """Ekstrak File ID dari berbagai format tautan Google Drive."""
    if not link or not str(link).strip():
        return None
    link = str(link).strip()
    # /file/d/FILE_ID/view
    m = re.search(r"/file/d/([a-zA-Z0-9_-]+)", link)
    if m:
        return m.group(1)
    # open?id=FILE_ID  atau  uc?id=FILE_ID
    m = re.search(r"[?&]id=([a-zA-Z0-9_-]+)", link)
    if m:
        return m.group(1)
    # File ID polos
    if re.fullmatch(r"[a-zA-Z0-9_-]{10,}", link):
        return link
    return None


def _download_drive_image(file_id: str):
    """
    Unduh satu file gambar dari Google Drive (public/shared link) via requests.
    Mengembalikan (BytesIO PNG, PIL.Image).
    Mendukung HEIC/HEIF jika pillow-heif terpasang.
    """
    url = f"https://drive.google.com/uc?export=download&id={file_id}"
    session = requests.Session()

    for attempt in range(1, _MAX_RETRIES + 1):
        try:
            resp = session.get(url, stream=True, timeout=30)
            # Tangani halaman konfirmasi Google Drive (file besar)
            if "text/html" in resp.headers.get("Content-Type", ""):
                token = None
                for k, v in resp.cookies.items():
                    if k.startswith("download_warning"):
                        token = v
                        break
                if token:
                    resp = session.get(url + f"&confirm={token}",
                                       stream=True, timeout=30)
                else:
                    resp = session.get(
                        f"https://drive.usercontent.google.com/download"
                        f"?id={file_id}&export=download&confirm=t",
                        stream=True, timeout=30,
                    )
            resp.raise_for_status()
            raw_bytes = resp.content
            break
        except requests.RequestException as exc:
            if attempt == _MAX_RETRIES:
                raise RuntimeError(
                    f"Gagal mengunduh file {file_id} setelah "
                    f"{_MAX_RETRIES} percobaan: {exc}"
                ) from exc
            time.sleep(_RETRY_DELAY_S)

    img = Image.open(io.BytesIO(raw_bytes))
    img.load()
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    png_fh = io.BytesIO()
    img.save(png_fh, format="PNG")
    png_fh.seek(0)
    return png_fh, img


# ── Helper OOXML / layout ──────────────────────────────────────────────────

def _set_cell_width(cell, width_cm: float):
    """Paksa lebar sel tabel ke width_cm."""
    cell.width = Cm(width_cm)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(width_cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _clear_cell_borders(cell):
    """Paksa semua sisi border sel menjadi 'nil' (tidak ada garis)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _remove_table_borders(table):
    """Hapus border tabel secara keseluruhan."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)


def _lock_row_height(row, height_cm: float):
    """Kunci tinggi baris (hRule=exact) dan larang baris terpotong lintas halaman."""
    trPr = row._tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)
    cant_split = OxmlElement("w:cantSplit")
    trPr.append(cant_split)


def _tighten_paragraph(paragraph):
    """Hilangkan spacing before/after & line-spacing berlebih bawaan Word."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.line_spacing = 1.0


def _add_centered_text(paragraph, text: str, bold: bool = False, size: int = 10):
    """Tambah run teks rata tengah ke paragraph."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return paragraph


def _fit_box(img_w: int, img_h: int, box_w: float, box_h: float):
    """Hitung dimensi agar gambar muat di dalam kotak, mempertahankan rasio aspek."""
    aspect = img_w / img_h
    target_w = box_w
    target_h = box_w / aspect
    if target_h > box_h:
        target_h = box_h
        target_w = box_h * aspect
    return target_w, target_h


def _fill_person_cell(cell, row_data: dict, idx: int) -> list:
    """
    Isi satu sel tabel grid 2x2 untuk satu orang.
    Mengembalikan daftar string peringatan (jika ada).
    """
    warnings_out = []
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _clear_cell_borders(cell)

    nik      = str(row_data.get("nik", "-")).strip() or "-"
    nama     = str(row_data.get("nama_lengkap", "-")).strip() or "-"
    wilayah  = str(row_data.get("wilayah_tugas", "-")).strip() or "-"
    operator = str(row_data.get("operator", "-")).strip() or "-"
    no_telp  = str(row_data.get("no_telp", "-")).strip() or "-"
    link     = str(row_data.get("link_bukti_terima", "")).strip()

    # Baris 1: Nama
    if cell.paragraphs:
        p_nama = cell.paragraphs[0]
        p_nama.text = ""
    else:
        p_nama = cell.add_paragraph()
    _add_centered_text(p_nama, nama, bold=True, size=12)

    # Baris 2: Foto
    p_img = cell.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(p_img)
    run_img = p_img.add_run()

    file_id = _extract_file_id(link)
    if file_id:
        if not HAS_PIL:
            run_img.text = "[Pillow tidak tersedia — foto dilewati]"
            run_img.italic = True
            warnings_out.append(
                f"[{idx + 1}] Pillow tidak terpasang; foto dilewati."
            )
        else:
            try:
                fh, img = _download_drive_image(file_id)
                img_w, img_h = img.size
                target_w, target_h = _fit_box(
                    img_w, img_h,
                    IMAGE_BOX_WIDTH_CM,
                    IMAGE_BOX_HEIGHT_CM,
                )
                fh.seek(0)
                run_img.add_picture(
                    fh, width=Cm(target_w), height=Cm(target_h)
                )
            except Exception as exc:
                msg = str(exc)
                if "403" in msg or "forbidden" in msg.lower():
                    reason = f"Akses ditolak (403) untuk file {file_id}"
                elif "404" in msg:
                    reason = f"File {file_id} tidak ditemukan (404)"
                else:
                    reason = f"Gagal memuat file {file_id}: {msg}"
                run_img.text = "[Gagal memuat gambar]"
                run_img.italic = True
                warnings_out.append(f"[{idx + 1}] {reason}")
    else:
        run_img.text = "[Foto tidak tersedia]"
        run_img.italic = True
        if link:
            warnings_out.append(
                f"[{idx + 1}] Tautan tidak dikenali: {link[:80]}"
            )
        else:
            warnings_out.append(
                f"[{idx + 1}] Kolom link_bukti_terima kosong."
            )

    # Baris 3: Metadata detail
    detail_lines = [
        f"NIK: {nik}",
        f"Wilayah Tugas: {wilayah}",
        f"Operator: {operator}",
        f"No Telp: {no_telp}",
    ]
    p_detail = cell.add_paragraph()
    p_detail.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _tighten_paragraph(p_detail)
    for i, line in enumerate(detail_lines):
        r = p_detail.add_run(line)
        r.font.size = Pt(9)
        if i < len(detail_lines) - 1:
            r.add_break()

    return warnings_out


# ── Validasi & template input ──────────────────────────────────────────────

def validate_input(file_path: str):
    """
    Validasi struktur file Excel input Bukti Terima.
    Returns (is_valid, errors, dfs) — dfs = {sheet_name: DataFrame}.
    """
    errors, dfs = [], {}
    try:
        xl = pd.ExcelFile(file_path)
    except Exception as exc:
        return False, [f"Gagal membaca file Excel: {exc}"], {}

    if SHEET_NAME not in xl.sheet_names:
        errors.append(
            f"Sheet '{SHEET_NAME}' tidak ditemukan. "
            f"Sheet yang tersedia: {', '.join(xl.sheet_names)}"
        )
        return False, errors, {}

    df = xl.parse(SHEET_NAME, dtype=str)
    df.columns = [str(c).strip() for c in df.columns]
    df = df.fillna("")
    dfs[SHEET_NAME] = df

    missing = [c for c in REQUIRED_COLUMNS if c not in df.columns]
    if missing:
        errors.append(
            f"Sheet '{SHEET_NAME}' kehilangan kolom: {', '.join(missing)}"
        )
        return False, errors, dfs

    if len(df) == 0:
        errors.append(f"Sheet '{SHEET_NAME}' tidak memiliki baris data.")
        return False, errors, dfs

    return True, [], dfs


def generate_input_template(file_path: str):
    """Buat template Excel input Bukti Terima (dengan baris contoh)."""
    import openpyxl
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = SHEET_NAME

    header_font  = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    header_fill  = PatternFill(
        start_color="4F81BD", end_color="4F81BD", fill_type="solid"
    )
    centered     = Alignment(horizontal="center", vertical="center")
    left_align   = Alignment(horizontal="left",   vertical="center")
    thin_border  = Border(
        left   = Side(style="thin", color="BFBFBF"),
        right  = Side(style="thin", color="BFBFBF"),
        top    = Side(style="thin", color="BFBFBF"),
        bottom = Side(style="thin", color="BFBFBF"),
    )

    all_cols = ["timestamp"] + REQUIRED_COLUMNS
    for col_idx, col_name in enumerate(all_cols, 1):
        cell = ws.cell(row=1, column=col_idx, value=col_name)
        cell.font      = header_font
        cell.fill      = header_fill
        cell.alignment = centered
        cell.border    = thin_border

    sample = {
        "timestamp":         "2026-07-29 15:51:27",
        "nik":               "6304XXXXXXXXXXXX",
        "nama_lengkap":      "Nama Petugas",
        "wilayah_tugas":     "[050] ANJIR MUARA",
        "operator":          "BY.U",
        "no_telp":           "08xxxxxxxxxx",
        "link_bukti_terima": "https://drive.google.com/open?id=FILE_ID_DISINI",
    }
    for col_idx, col_name in enumerate(all_cols, 1):
        cell = ws.cell(row=2, column=col_idx, value=sample.get(col_name, ""))
        cell.alignment = left_align
        cell.border    = thin_border

    for col in ws.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        letter  = openpyxl.utils.get_column_letter(col[0].column)
        ws.column_dimensions[letter].width = max(max_len + 3, 12)

    wb.save(file_path)


# ── Generator utama ────────────────────────────────────────────────────────

def iter_generate(dfs: dict, out_dir: str):
    """
    Generator utama Bukti Terima.

    Membuat satu file DOCX berisi semua petugas dalam grid 2x2 tetap per
    halaman A4 (tanpa border), foto diunduh satu per satu saat giliran
    orang itu diproses (lazy download).

    Yields event dict:
      {"t": "log",      "msg": str, "level": str}
      {"t": "progress", "done": int, "total": int}
      {"t": "file",     "path": str}
      {"t": "done",     "generated": list[str]}
    """
    if not HAS_DOCX:
        yield {"t": "log", "msg": "python-docx tidak tersedia.", "level": "ERROR"}
        yield {"t": "done", "generated": []}
        return

    df = dfs.get(SHEET_NAME)
    if df is None or df.empty:
        yield {"t": "log",
               "msg": f"Sheet '{SHEET_NAME}' kosong atau tidak ditemukan.",
               "level": "ERROR"}
        yield {"t": "done", "generated": []}
        return

    rows  = df.to_dict("records")
    total = len(rows)
    n_pages = (total + 3) // 4

    yield {"t": "log",
           "msg": f"Memproses {total} petugas → {n_pages} halaman (grid 2x2)…",
           "level": "INFO"}

    if not HAS_PIL:
        yield {"t": "log",
               "msg": "Pillow tidak terpasang — foto akan dilewati.",
               "level": "WARN"}

    # ── Buat dokumen kosong A4 ─────────────────────────────────────────────
    doc = Document()
    section = doc.sections[0] if doc.sections else doc.add_section()
    section.page_width    = Cm(PAGE_W_CM)
    section.page_height   = Cm(PAGE_H_CM)
    section.left_margin   = Cm(MARGIN_CM)
    section.right_margin  = Cm(MARGIN_CM)
    section.top_margin    = Cm(MARGIN_CM)
    section.bottom_margin = Cm(MARGIN_CM)

    # Hapus paragraf kosong awal (jika ada) agar tabel mulai persis di batas atas
    if doc.paragraphs:
        p0 = doc.paragraphs[0]
        p0_elem = p0._element
        parent = p0_elem.getparent()
        if parent is not None:
            parent.remove(p0_elem)

    # ── Proses per halaman (4 orang per halaman) ───────────────────────────
    def _chunk(lst, n):
        for i in range(0, len(lst), n):
            yield lst[i:i + n]

    done_count  = 0
    all_warnings = []

    for page_idx, group in enumerate(_chunk(rows, 4)):
        if page_idx > 0:
            doc.add_page_break()

        # Tabel 2 kolom x 2 baris tetap per halaman
        table = doc.add_table(rows=2, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit   = False
        table.style     = None  # hindari style Word yang membawa border

        _remove_table_borders(table)

        # Kunci tinggi tiap baris agar tidak meluber ke halaman berikutnya
        for table_row in table.rows:
            _lock_row_height(table_row, ROW_HEIGHT_CM)

        for i in range(4):
            r, c = divmod(i, 2)
            cell = table.cell(r, c)
            _set_cell_width(cell, COL_WIDTH_CM)

            if i < len(group):
                row_data = group[i]
                nama = str(row_data.get("nama_lengkap",
                                        f"Orang {done_count + 1}")).strip()
                yield {"t": "log",
                       "msg": f"[{done_count + 1}/{total}] Memproses: {nama}",
                       "level": "INFO"}

                warns = _fill_person_cell(cell, row_data, done_count)
                for w in warns:
                    all_warnings.append(w)
                    yield {"t": "log", "msg": f"  {w}", "level": "WARN"}

                done_count += 1
            else:
                # Sel kosong — halaman terakhir dengan < 4 orang
                if cell.paragraphs:
                    cell.paragraphs[0].text = ""
                _clear_cell_borders(cell)

        yield {"t": "progress", "done": done_count, "total": total}

    # ── Simpan DOCX ────────────────────────────────────────────────────────
    os.makedirs(out_dir, exist_ok=True)
    ts       = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out_name = f"Bukti_Terima_{ts}.docx"
    out_path = os.path.join(out_dir, out_name)

    try:
        doc.save(out_path)
    except Exception as exc:
        yield {"t": "log",
               "msg": f"Gagal menyimpan DOCX: {exc}",
               "level": "ERROR"}
        yield {"t": "done", "generated": []}
        return

    yield {"t": "log",
           "msg": f"DOCX tersimpan: {out_name}  "
                  f"({n_pages} halaman, grid 2x2 tetap, tanpa border)",
           "level": "OK"}

    if all_warnings:
        yield {"t": "log",
               "msg": f"Total peringatan unduhan foto: {len(all_warnings)}",
               "level": "WARN"}

    yield {"t": "file", "path": out_path}
    yield {"t": "done", "generated": [out_path]}

