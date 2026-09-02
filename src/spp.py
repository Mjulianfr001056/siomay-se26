"""
Generator Surat Pernyataan Penyelesaian (SPP) — Sensus Ekonomi 2026.

Port dari notebook 'generator/Generator_SPP_PPL.ipynb' dan
'generator/Generator_SPP_PML.ipynb', diadaptasi ke format input
dengan satu file Excel berisi 3 sheet.

Sheet input:
  - data_mitra      : nik, nama_lengkap, jabatan  (PPL & PML)
  - no_spk          : nik, no_spk, no_urut_spp_t1
  - alokasi_usaha   : nik_ppl, nik_pml, target, capaian, persentase

Fungsi inti ``iter_generate()`` adalah *generator* yang melempar
event log/progress/file agar UI bisa menampilkan proses secara langsung.
"""
import copy
import os
import re
import tempfile

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Row

# -- Skema sheet wajib ------------------------------------------------
SHEET_DATA_MITRA = "data_mitra"
SHEET_NO_SPK = "no_spk"
SHEET_ALOKASI = "alokasi_usaha"

SHEET_NAME = "data_mitra"

REQUIRED_SCHEMA = {
    SHEET_DATA_MITRA: ["nik", "nama_lengkap", "jabatan"],
    SHEET_NO_SPK: ["nik", "no_spk", "no_urut_spp_t1"],
    SHEET_ALOKASI: ["nik_ppl", "nik_pml", "target", "capaian", "persentase"],
}

# -- Column names -----------------------------------------------------
COL_NIK = "nik"
COL_NAMA = "nama_lengkap"
COL_JABATAN = "jabatan"
COL_NO_SPK = "no_spk"
COL_NO_URUT_SPP_T1 = "no_urut_spp_t1"
COL_TARGET = "target"
COL_CAPAIAN = "capaian"
COL_PERSENTASE = "persentase"
COL_NIK_PPL = "nik_ppl"
COL_NIK_PML = "nik_pml"

# -- PML table column widths (twips) ----------------------------------
LAMPIRAN_COL_WIDTHS = [700, 3200, 1500, 1800, 1500, 0]


# =====================================================================
#  Validation
# =====================================================================
def validate_input(file_path: str):
    """Validate the SPP input Excel file.

    Returns ``(ok, errors, dfs)`` where *dfs* maps sheet_name -> DataFrame.
    """
    errors = []
    dfs = {}
    try:
        xls = pd.ExcelFile(file_path)
    except Exception as exc:
        return False, [f"Gagal membaca file Excel: {exc}"], {}

    for sheet_name, required_cols in REQUIRED_SCHEMA.items():
        if sheet_name not in xls.sheet_names:
            errors.append(f"Sheet '{sheet_name}' tidak ditemukan.")
            continue
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)
        df.columns = df.columns.str.strip()
        df = df.fillna("")
        missing = [c for c in required_cols if c not in df.columns]
        if missing:
            errors.append(
                f"Sheet '{sheet_name}' kekurangan kolom: " + ", ".join(missing)
            )
        else:
            dfs[sheet_name] = df

    ok = len(errors) == 0
    return ok, errors, dfs


# =====================================================================
#  Helper functions (ported from notebook)
# =====================================================================
def _format_persentase(raw):
    """Format percentage to Indonesian locale (comma as decimal, 2 dp)."""
    try:
        return f"{float(raw):.2f}".replace(".", ",")
    except (ValueError, TypeError):
        return str(raw).strip()


def _strip_placeholder_fmt(run):
    """Convert editorial placeholder formatting to normal black text."""
    rpr = run._r.find(qn("w:rPr"))
    if rpr is None:
        return
    color = rpr.find(qn("w:color"))
    if color is not None:
        val = color.get(qn("w:val")) or ""
        if val.lower() in ("0000ff", "ee0000", "ff0000"):
            color.set(qn("w:val"), "000000")
            color.attrib.pop(qn("w:themeColor"), None)
            color.attrib.pop(qn("w:themeTint"), None)
            color.attrib.pop(qn("w:themeShade"), None)
    u = rpr.find(qn("w:u"))
    if u is not None:
        rpr.remove(u)


def replace_text_preserving_runs(doc: Document, replacements: dict) -> None:
    """Replace ``{{key}}`` placeholders while preserving run formatting."""

    def _process_paragraphs(paragraphs):
        for para in paragraphs:
            runs = para.runs
            if not runs:
                continue
            full_text = ""
            char_map = []
            for ri, run in enumerate(runs):
                for ci, ch in enumerate(run.text):
                    char_map.append((ri, ci))
                    full_text += ch
            matches = list(re.finditer(r"\{\{[^}]+\}\}", full_text))
            if not matches:
                continue
            for m in reversed(matches):
                ph = m.group(0)
                if ph not in replacements:
                    continue
                val = str(replacements[ph])
                start, end = m.start(), m.end()
                touched = []
                for ci in range(start, min(end, len(char_map))):
                    ri = char_map[ci][0]
                    if not touched or touched[-1] != ri:
                        touched.append(ri)
                if not touched:
                    continue
                first_ri = touched[0]
                last_ri = touched[-1]
                first_run_start = next(
                    i for i, (ri, _) in enumerate(char_map) if ri == first_ri
                )
                prefix = full_text[first_run_start:start]
                last_run_end = (
                    max(i for i, (ri, _) in enumerate(char_map) if ri == last_ri)
                    + 1
                )
                suffix = full_text[end:last_run_end] if end < last_run_end else ""
                runs[first_ri].text = prefix + val
                _strip_placeholder_fmt(runs[first_ri])
                for ri in touched[1:]:
                    runs[ri].text = suffix if (ri == last_ri and suffix) else ""

    _process_paragraphs(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                _process_paragraphs(cell.paragraphs)


def _set_cell_text(cell, text):
    """Write text to a table cell preserving existing formatting."""
    para = cell.paragraphs[0]
    if para.runs:
        ref_run = para.runs[0]
        para.clear()
        new_run = para.add_run(str(text))
        rpr_ref = ref_run._r.find(qn("w:rPr"))
        if rpr_ref is not None:
            rpr_new = copy.deepcopy(rpr_ref)
            color = rpr_new.find(qn("w:color"))
            if color is not None:
                val = color.get(qn("w:val")) or ""
                if val.lower() in ("0000ff", "ff0000"):
                    rpr_new.remove(color)
            u = rpr_new.find(qn("w:u"))
            if u is not None:
                rpr_new.remove(u)
            new_run._r.insert(0, rpr_new)
    else:
        para.text = str(text)


def _fix_lampiran_layout(table):
    """Center the PML table, fix its widths, and center-align data cells."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is not None:
        jc = tbl_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tbl_pr.append(jc)
        jc.set(qn("w:val"), "center")
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        if tbl_ind is not None:
            tbl_pr.remove(tbl_ind)
        tbl_w = tbl_pr.find(qn("w:tblW"))
        if tbl_w is not None:
            tbl_w.set(qn("w:w"), "0")
            tbl_w.set(qn("w:type"), "auto")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is not None:
        cols = grid.findall(qn("w:gridCol"))
        for col, w in zip(cols, LAMPIRAN_COL_WIDTHS):
            col.set(qn("w:w"), str(w))
        tgc = grid.find(qn("w:tblGridChange"))
        if tgc is not None:
            grid.remove(tgc)

    for row in table.rows:
        for cell, w in zip(row.cells, LAMPIRAN_COL_WIDTHS):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(w))
            p = cell.paragraphs[0]
            pPr = p._p.get_or_add_pPr()
            pJc = OxmlElement("w:jc")
            pJc.set(qn("w:val"), "center")
            old_jc = pPr.find(qn("w:jc"))
            if old_jc is not None:
                pPr.remove(old_jc)
            pPr.append(pJc)


def _fill_pml_lampiran_table(table, ppl_rows):
    """Fill PML lampiran table with PPL rows + Jumlah summary row."""
    template_row_tr = copy.deepcopy(table.rows[2]._tr)

    # Delete 3 example data rows (index 2, 3, 4)
    for r in list(table.rows)[2:5]:
        r._tr.getparent().remove(r._tr)

    # After deletion, 'Jumlah' is now at index 2
    jumlah_tr = table.rows[2]._tr
    jumlah_target = 0.0
    jumlah_capaian = 0.0

    for i, ppl in enumerate(ppl_rows, start=1):
        new_tr = copy.deepcopy(template_row_tr)
        jumlah_tr.addprevious(new_tr)
        new_row = _Row(new_tr, table)
        cells = new_row.cells
        _set_cell_text(cells[0], f"{i}.")
        _set_cell_text(cells[1], ppl["nama"])
        _set_cell_text(cells[2], f"{ppl['target']:.0f}")
        _set_cell_text(cells[3], f"{ppl['capaian']:.0f}")
        _set_cell_text(cells[4], ppl["persentase_str"])
        jumlah_target += ppl["target"]
        jumlah_capaian += ppl["capaian"]

    jumlah_row = _Row(jumlah_tr, table)
    jc_cells = jumlah_row.cells
    jumlah_persentase = (
        (jumlah_capaian / jumlah_target * 100) if jumlah_target else 0
    )
    _set_cell_text(jc_cells[2], f"{jumlah_target:.0f}")
    _set_cell_text(jc_cells[3], f"{jumlah_capaian:.0f}")
    _set_cell_text(jc_cells[4], _format_persentase(jumlah_persentase))
    _fix_lampiran_layout(table)


# =====================================================================
#  Generator
# =====================================================================
def iter_generate(
        kind, dfs, template_path, out_dir,
        number_column=COL_NO_URUT_SPP_T1,
        number_placeholder="no_urut_spp_t1", termin_label=""):
    """Generate SPP documents.

    Yields event dicts::

        {"t": "log",      "msg": str, "level": str}
        {"t": "progress", "done": int, "total": int}
        {"t": "file",     "path": str}
        {"t": "done",     "generated": list[str]}

    Parameters
    ----------
    kind : ``"ppl"`` | ``"pml"``
    dfs : dict  -- sheet_name -> DataFrame (from validate_input)
    template_path : str  -- path to .docx template
    out_dir : str  -- output directory
    """
    os.makedirs(out_dir, exist_ok=True)
    generated = []

    df_mitra = dfs.get(SHEET_DATA_MITRA, pd.DataFrame())
    df_spk = dfs.get(SHEET_NO_SPK, pd.DataFrame())
    df_alokasi = dfs.get(SHEET_ALOKASI, pd.DataFrame())

    if df_mitra.empty:
        yield {"t": "log", "msg": "Sheet 'data_mitra' kosong.", "level": "ERROR"}
        yield {"t": "done", "generated": generated}
        return

    # Merge mitra + spk on NIK
    df_mitra[COL_NIK] = df_mitra[COL_NIK].str.strip()
    df_spk[COL_NIK] = df_spk[COL_NIK].str.strip()
    df_merged = pd.merge(df_mitra, df_spk, on=COL_NIK, how="left")
    df_merged = df_merged.fillna("")

    # Filter by kind
    if kind == "ppl":
        df_officers = df_merged[
            df_merged[COL_JABATAN].str.strip().str.upper() == "PPL"
        ].copy()
        label = "PPL"
    else:
        df_officers = df_merged[
            df_merged[COL_JABATAN].str.strip().str.upper() == "PML"
        ].copy()
        label = "PML"

    total = len(df_officers)
    yield {"t": "log", "msg": f"Memproses {total} {label}...", "level": "STEP"}

    if total == 0:
        yield {
            "t": "log",
            "msg": f"Tidak ditemukan {label} di sheet '{SHEET_DATA_MITRA}'.",
            "level": "ERROR",
        }
        yield {"t": "done", "generated": generated}
        return

    # Build name lookup from mitra for resolving PPL names in PML table
    name_lookup = {}
    for _, mrow in df_mitra.iterrows():
        n = str(mrow.get(COL_NAMA, "")).strip()
        k = str(mrow.get(COL_NIK, "")).strip()
        if k and n:
            name_lookup[k] = n

    for idx, (_, row) in enumerate(df_officers.iterrows()):
        nik = str(row.get(COL_NIK, "")).strip()
        nama = str(row.get(COL_NAMA, "")).strip()
        no_spk_val = str(row.get(COL_NO_SPK, "")).strip()
        no_input = str(row.get(number_column, "")).strip()

        if not nama or nama.lower() == "nan":
            yield {
                "t": "log",
                "msg": f"  [{idx+1}] Skip -- nama kosong.",
                "level": "WARN",
            }
            continue

        yield {
            "t": "log",
            "msg": f"  [{idx+1}/{total}] {nama} | no_input={no_input}",
            "level": "INFO",
        }

        if kind == "ppl":
            file_path = _generate_ppl_doc(
                nik, nama, no_spk_val, no_input, df_alokasi, template_path,
                number_placeholder,
            )
        else:
            file_path = _generate_pml_doc(
                nik, nama, no_spk_val, no_input, df_alokasi,
                name_lookup, template_path, number_placeholder,
            )

        if file_path:
            safe_name = re.sub(r"[^A-Za-z0-9]+", "_", nama)
            out_name = (
                f"SPP_{label.upper()}{termin_label}_{idx+1:03d}_{safe_name}.docx"
            )
            dest = os.path.join(out_dir, out_name)
            os.replace(file_path, dest)
            generated.append(dest)
            yield {"t": "file", "path": dest}
            yield {"t": "log", "msg": f"    Tersimpan: {out_name}", "level": "OK"}
        else:
            yield {
                "t": "log",
                "msg": f"    Peringatan: tidak ada data alokasi untuk {nama} — dilewati.",
                "level": "WARN",
            }

        yield {"t": "progress", "done": idx + 1, "total": total}

    yield {"t": "done", "generated": generated}


def _generate_ppl_doc(
        nik, nama, no_spk_val, no_input, df_alokasi, template_path,
        number_placeholder="no_urut_spp_t1"):
    """Generate a single SPP PPL document. Returns output path or None."""
    df_alokasi[COL_NIK_PPL] = df_alokasi[COL_NIK_PPL].str.strip()

    grp = df_alokasi[df_alokasi[COL_NIK_PPL] == nik]
    total_target = 0.0
    total_capaian = 0.0
    for _, arow in grp.iterrows():
        try:
            total_target += float(arow[COL_TARGET])
        except (ValueError, TypeError):
            pass
        try:
            total_capaian += float(arow[COL_CAPAIAN])
        except (ValueError, TypeError):
            pass

    jml_usaha = f"{total_target:.0f}" if total_target else ""
    jml_usaha_min = f"{total_capaian:.0f}" if total_capaian else ""
    persentase = _format_persentase(
        (total_capaian / total_target * 100) if total_target else 0
    )

    doc = Document(template_path)
    replace_text_preserving_runs(
        doc,
        {
            "{{" + number_placeholder + "}}": no_input,
            "{{nama_lengkap}}": nama,
            "{{nik}}": nik,
            "{{no_spk}}": no_spk_val,
            "{{jml_usaha}}": jml_usaha,
            "{{jml_usaha_min}}": jml_usaha_min,
            "{{persentase}}": persentase,
        },
    )

    tmp_path = tempfile.mktemp(suffix=".docx", prefix="spp_ppl_")
    doc.save(tmp_path)
    return tmp_path


def _generate_pml_doc(
        nik, nama, no_spk_val, no_input, df_alokasi, name_lookup,
        template_path, number_placeholder="no_urut_spp_t1"):
    """Generate a single SPP PML document. Returns output path or None."""
    df_alokasi[COL_NIK_PPL] = df_alokasi[COL_NIK_PPL].str.strip()
    df_alokasi[COL_NIK_PML] = df_alokasi[COL_NIK_PML].str.strip()

    grp = df_alokasi[df_alokasi[COL_NIK_PML] == nik].copy()

    if grp.empty:
        return None

    ppl_rows = []
    for _, prow in grp.iterrows():
        try:
            target = float(prow[COL_TARGET])
        except (ValueError, TypeError):
            target = 0.0
        try:
            capaian = float(prow[COL_CAPAIAN])
        except (ValueError, TypeError):
            capaian = 0.0

        ppl_nik = str(prow[COL_NIK_PPL]).strip()
        ppl_name = name_lookup.get(ppl_nik, ppl_nik)

        ppl_rows.append({
            "nama": ppl_name,
            "target": target,
            "capaian": capaian,
            "persentase_str": _format_persentase(
                (capaian / target * 100) if target else 0
            ),
        })

    doc = Document(template_path)
    replace_text_preserving_runs(
        doc,
        {
            "{{" + number_placeholder + "}}": no_input,
            "{{nama_lengkap}}": nama,
            "{{nik}}": nik,
            "{{no_spk}}": no_spk_val,
        },
    )

    # Fill lampiran table (table index 1 in the template)
    if len(doc.tables) > 1:
        _fill_pml_lampiran_table(doc.tables[1], ppl_rows)

    tmp_path = tempfile.mktemp(suffix=".docx", prefix="spp_pml_")
    doc.save(tmp_path)
    return tmp_path
