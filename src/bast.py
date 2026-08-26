"""
Generator Berita Acara Serah Terima (BAST) — Sensus Ekonomi 2026.

Port dari alur yang tersirat dari template BAST PPL dan BAST PML.

Sheet input (shared file untuk PPL & PML):
  - data_mitra       : nik, nama_mitra, jabatan (ppl|pml),
                        no_spk, no_urut_bast, bukti_dukung_bast
  - alokasi_supervisi: nik_ppl, nik_pml
  - alokasi_tugas    : nik_ppl, nik_pml, idsubsls, nmsls, kdkec, kddesa,
                        kdsls, kdsubsls, nmdesa
  - kode_kecamatan   : kode, kecamatan  (opsional)

Template DOCX placeholders:
  {{no_urut_bast}}, {{no_spk}}, {{nama_lengkap}}, {{nik}},
  {{jml_sls}}, {{bukti_dukung_bast}}

Template Tables:
  Table 0 – Tanda tangan (PIHAK KEDUA / PIHAK PERTAMA)
  Table 1 – Uraian pekerjaan ({{jml_sls}})
  Table 2 – Daftar wilayah kerja (data dynamic)
  Table 3 – Tanda tangan penutup
"""
import os

import pandas as pd
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Reuse image and placeholder helpers from BAPP PPL
from src.bapp_ppl import (
    replace_text_preserving_runs,
    insert_gdrive_images,
    _slug,
)

# -- Skema sheet wajib ------------------------------------------------
REQUIRED_SCHEMA = {
    "data_mitra": [
        "nik", "nama_mitra", "jabatan", "no_spk",
        "no_urut_bast", "bukti_dukung_bast",
    ],
    "alokasi_supervisi": ["nik_ppl", "nik_pml"],
    "alokasi_tugas": [
        "nik_ppl", "nik_pml", "idsubsls", "nmsls",
        "kdkec", "kddesa", "kdsls", "kdsubsls", "nmdesa",
    ],
}
OPTIONAL_SCHEMA = {
    "kode_kecamatan": ["kode", "kecamatan"],
}

SHEET_NAME = "data_mitra"
BUKTI_PLACEHOLDER = "{{bukti_dukung_bast}}"


def _norm(v) -> str:
    """Normalisasi nilai sel menjadi string bersih."""
    if pd.isna(v):
        return ""
    return str(v).strip()



# =====================================================================
#  Validation
# =====================================================================
def validate_input(file_path: str):
    """Validate the BAST input Excel file.

    Returns ``(ok, errors, dfs)`` where *dfs* maps sheet_name -> DataFrame.
    """
    errors, dfs = [], {}
    try:
        import openpyxl  # noqa: F401
        xls = pd.ExcelFile(file_path)
    except Exception as e:
        return False, [f"Gagal membaca file Excel: {e}"], {}

    for sheet_name, required_cols in {**REQUIRED_SCHEMA, **OPTIONAL_SCHEMA}.items():
        if sheet_name not in xls.sheet_names:
            if sheet_name in REQUIRED_SCHEMA:
                errors.append(
                    f"Sheet '{sheet_name}' tidak ditemukan. "
                    f"Sheet yang tersedia: {', '.join(xls.sheet_names)}"
                )
            continue
        df = pd.read_excel(xls, sheet_name=sheet_name, dtype=str)
        df.columns = df.columns.str.strip()
        df = df.fillna("")
        dfs[sheet_name] = df
        if sheet_name in REQUIRED_SCHEMA:
            missing = [c for c in required_cols if c not in df.columns]
            if missing:
                errors.append(
                    f"Sheet '{sheet_name}' kekurangan kolom: "
                    + ", ".join(missing)
                )

    if not errors:
        for sheet_name in REQUIRED_SCHEMA:
            if sheet_name in dfs and len(dfs[sheet_name]) == 0:
                errors.append(f"Sheet '{sheet_name}' kosong.")

    return (len(errors) == 0), errors, dfs


# =====================================================================
#  Table helpers (adapted from lampiran_spk.py)
# =====================================================================
def _make_cell(text, ref_cell, vmerge=None):
    """Buat elemen <w:tc> baru berdasarkan ref_cell."""
    nc = deepcopy(ref_cell._tc)
    if vmerge is not None:
        tcPr = nc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            nc.insert(0, tcPr)
        for old in tcPr.findall(qn("w:vMerge")):
            tcPr.remove(old)
        vm = OxmlElement("w:vMerge")
        if vmerge == "restart":
            vm.set(qn("w:val"), "restart")
        tcPr.append(vm)
    for p in nc.findall(qn("w:p")):
        for r in p.findall(qn("w:r")):
            p.remove(r)
        if vmerge == "continue":
            continue
        r_el = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        r_el.append(rpr)
        t_el = OxmlElement("w:t")
        t_el.text = str(text)
        t_el.set(
            "{http://www.w3.org/XML/1998/namespace}space", "preserve"
        )
        r_el.append(t_el)
        p.append(r_el)
    return nc


def _append_row(tbl_xml, ref_row, values, vmerges=None):
    """Tambah baris baru di akhir tabel."""
    new_tr = OxmlElement("w:tr")
    ref_trPr = ref_row._tr.find(qn("w:trPr"))
    if ref_trPr is not None:
        new_tr.append(deepcopy(ref_trPr))
    rc = ref_row.cells
    for idx, val in enumerate(values):
        vm = vmerges[idx] if vmerges else None
        new_tr.append(_make_cell(val, rc[idx], vmerge=vm))
    tbl_xml.append(new_tr)


def _detect_data_start(table):
    """Baris data pertama = setelah semua baris header."""
    data_start = 0
    markers = (
        "(1)", "(2)", "(3)", "(4)", "(5)",
        "No", "KECAMATAN/DISTRIK",
        "Nama Petugas Lapangan Sensus",
    )
    for i, row in enumerate(table.rows):
        texts = [c.text.strip() for c in row.cells]
        if any(t in markers for t in texts):
            data_start = i + 1
    return data_start


def _find_wilayah_table(doc, need_nama_col=False):
    """Cari tabel DAFTAR WILAYAH KERJA berdasarkan header."""
    for t in doc.tables:
        if t.rows:
            header = " ".join(c.text for c in t.rows[0].cells).lower()
            ok = "kecamatan" in header
            if need_nama_col:
                ok = ok and "nama" in header
            if ok:
                return t
    # Fallback: tabel ketiga (index 2) — tabel wilayah
    if len(doc.tables) >= 3:
        return doc.tables[2]
    return None


def _kec_text(kdkec, kec_map):
    """Format kolom kecamatan: '[kode] Nama Kecamatan'."""
    name = kec_map.get(str(kdkec).strip(), "")
    return f"[{kdkec}] {name}" if name else (
        f"[{kdkec}]" if kdkec else ""
    )


def _desa_text(kddesa, nmdesa):
    """Format kolom desa: '[kode] Nama Desa'."""
    if kddesa and nmdesa:
        return f"[{kddesa}] {nmdesa}"
    return nmdesa or ""


# =====================================================================
#  Wilayah population
# =====================================================================
def _populate_wilayah_ppl(table, wilayah_rows, log=None):
    """Isi tabel DAFTAR WILAYAH KERJA BAST PPL (4 kolom).

    wilayah_rows: list of (no_str, kec_text, desa_text, jml_sls_str)
    """
    tbl_xml = table._tbl
    all_trs = tbl_xml.findall(qn("w:tr"))
    data_start = _detect_data_start(table)
    for tr in all_trs[data_start:]:
        tbl_xml.remove(tr)
    if not wilayah_rows:
        if log:
            log("Tidak ada data wilayah kerja untuk PPL.", "WARN")
        return
    ref_row = table.rows[min(data_start, len(table.rows) - 1)]
    for (no, kec, desa, jml) in wilayah_rows:
        _append_row(tbl_xml, ref_row, [no, kec, desa, jml])


def _populate_wilayah_pml(table, pml_rows, log=None):
    """Isi tabel DAFTAR WILAYAH KERJA BAST PML (5 kolom, dengan nama PPL).

    pml_rows: list of (no_str, nama_ppl, kec_text, desa_text, jml_sls_str)
    nama_ppl hanya terisi pada baris pertama tiap kelompok PPL.
    """
    tbl_xml = table._tbl
    all_trs = tbl_xml.findall(qn("w:tr"))
    data_start = _detect_data_start(table)
    for tr in all_trs[data_start:]:
        tbl_xml.remove(tr)
    if not pml_rows:
        if log:
            log("Tidak ada data wilayah kerja untuk PML.", "WARN")
        return
    ref_row = table.rows[min(data_start, len(table.rows) - 1)]

    # Build rows with vMerge for PPL name grouping
    i = 0
    while i < len(pml_rows):
        no, nama, kec, desa, jml = pml_rows[i]
        span = 1
        j = i + 1
        while j < len(pml_rows) and pml_rows[j][1] == "":
            span += 1
            j += 1
        for k in range(span):
            no2 = pml_rows[i + k][0]
            nama2 = pml_rows[i + k][1] if k == 0 else ""
            kec2 = pml_rows[i + k][2]
            desa2 = pml_rows[i + k][3]
            jml2 = pml_rows[i + k][4]
            if k == 0 and span > 1:
                vm = "restart"
            elif k > 0:
                vm = "continue"
            else:
                vm = None
            _append_row(
                tbl_xml, ref_row,
                [no2, nama2, kec2, desa2, jml2],
                vmerges=[vm, vm, None, None, None],
            )
        i += span


# =====================================================================
#  Data building helpers
# =====================================================================
def _build_wilayah_data_ppl(nik, df_tugas, kec_map):
    """Susun data wilayah kerja untuk seorang PPL.

    Returns (jml_sls_total, rows_for_table).
    rows_for_table: list of (no, kec_text, desa_text, jml_sls)
    """
    grp = df_tugas[df_tugas["nik_ppl"] == nik].copy()
    if grp.empty:
        return 0, []

    jml_sls = (
        int(grp["idsubsls"].nunique())
        if "idsubsls" in grp.columns
        else len(grp)
    )

    grp_cols = ["kdkec", "kddesa", "nmdesa"]
    if "idsubsls" in grp.columns:
        grouped = (
            grp.groupby(grp_cols, sort=True)["idsubsls"]
            .nunique()
            .reset_index(name="jml_subsls")
        )
    else:
        grouped = (
            grp.groupby(grp_cols, sort=True)
            .size()
            .reset_index(name="jml_subsls")
        )

    rows = []
    for i, r in enumerate(grouped.itertuples(), start=1):
        rows.append((
            f"{i}.",
            _kec_text(r.kdkec, kec_map),
            _desa_text(r.kddesa, r.nmdesa),
            str(r.jml_subsls),
        ))
    return jml_sls, rows


def _build_wilayah_data_pml(nik_pml, df_tugas, df_mitra, kec_map):
    """Susun data wilayah kerja untuk seorang PML.

    Returns (jml_sls_total, rows_for_table).
    rows_for_table: list of (no, nama_ppl, kec_text, desa_text, jml_sls)
    """
    grp = df_tugas[df_tugas["nik_pml"] == nik_pml].copy()
    if grp.empty:
        return 0, []

    jml_sls = (
        int(grp["idsubsls"].nunique())
        if "idsubsls" in grp.columns
        else len(grp)
    )

    name_lookup = {}
    for _, r in df_mitra.iterrows():
        n = _norm(r.get("nik", ""))
        nm = _norm(r.get("nama_mitra", ""))
        if n and nm:
            name_lookup[n] = nm

    grp_cols = ["nik_ppl", "kdkec", "kddesa", "nmdesa"]
    if "idsubsls" in grp.columns:
        grouped = (
            grp.groupby(grp_cols, sort=True)["idsubsls"]
            .nunique()
            .reset_index(name="jml_subsls")
        )
    else:
        grouped = (
            grp.groupby(grp_cols, sort=True)
            .size()
            .reset_index(name="jml_subsls")
        )

    rows = []
    row_no = 1
    prev_ppl = None
    for r in grouped.itertuples():
        nik_p = _norm(r.nik_ppl)
        nama_ppl = name_lookup.get(nik_p, nik_p)
        display_name = nama_ppl if nama_ppl != prev_ppl else ""
        prev_ppl = nama_ppl
        rows.append((
            f"{row_no}." if display_name else "",
            display_name,
            _kec_text(r.kdkec, kec_map),
            _desa_text(r.kddesa, r.nmdesa),
            str(r.jml_subsls),
        ))
        if display_name:
            row_no += 1

    return jml_sls, rows


# =====================================================================
#  Document generation
# =====================================================================
def _generate_one_doc(row, kind, df_tugas, df_mitra, kec_map,
                      template_path, out_dir, idx, total):
    """Generate satu dokumen BAST (PPL atau PML). Yield event dicts."""
    nik = _norm(row.get("nik", ""))
    nama = _norm(row.get("nama_mitra", ""))
    no_spk_val = _norm(row.get("no_spk", ""))
    no_urut = _norm(row.get("no_urut_bast", ""))
    bukti_link = _norm(row.get("bukti_dukung_bast", ""))

    who = nama or nik
    yield {"t": "log", "level": "INFO",
           "msg": f"[{idx + 1}/{total}] {who} (NIK {nik})"}

    if kind == "ppl":
        jml_sls, wilayah_rows = _build_wilayah_data_ppl(
            nik, df_tugas, kec_map
        )
    else:
        jml_sls, wilayah_rows = _build_wilayah_data_pml(
            nik, df_tugas, df_mitra, kec_map
        )

    yield {"t": "log", "level": "INFO",
           "msg": f"   jml_sls={jml_sls}, "
                  f"wilayah={len(wilayah_rows)} baris"}

    doc = Document(template_path)
    replacements = {
        "{{no_urut_bast}}": no_urut,
        "{{no_spk}}": no_spk_val,
        "{{nama_lengkap}}": nama,
        "{{nik}}": nik,
        "{{jml_sls}}": str(jml_sls),
    }
    replace_text_preserving_runs(doc, replacements)

    # Populate wilayah kerja table (Table 2)
    need_nama = (kind == "pml")
    table = _find_wilayah_table(doc, need_nama_col=need_nama)
    if table is not None:
        if kind == "ppl":
            _populate_wilayah_ppl(table, wilayah_rows)
        else:
            _populate_wilayah_pml(table, wilayah_rows)
        yield {"t": "log", "level": "INFO",
               "msg": f"   Tabel wilayah diisi ({len(wilayah_rows)} baris)"}
    else:
        yield {"t": "log", "level": "WARN",
               "msg": "   Tabel wilayah kerja tidak ditemukan!"}

    # Insert screenshots bukti dukung
    if bukti_link:
        n_img, img_warnings = insert_gdrive_images(
            doc, bukti_link, placeholder=BUKTI_PLACEHOLDER
        )
        if n_img:
            yield {"t": "log", "level": "INFO",
                   "msg": f"   {n_img} screenshot disisipkan"}
        for w in img_warnings:
            yield {"t": "log", "level": "WARN", "msg": f"   {w}"}
    else:
        for p in doc.paragraphs:
            if BUKTI_PLACEHOLDER in p.text:
                for run in p.runs:
                    if BUKTI_PLACEHOLDER in run.text:
                        run.text = run.text.replace(BUKTI_PLACEHOLDER, "")

    kind_label = "PPL" if kind == "ppl" else "PML"
    safe_name = _slug(who)
    out_name = f"BAST_{kind_label}_{idx + 1:03d}_{nik}_{safe_name}.docx"
    out_path = os.path.join(out_dir, out_name)
    doc.save(out_path)
    yield {"t": "file", "path": out_path}
    yield {"t": "log", "level": "OK",
           "msg": f"   Tersimpan: {out_name}"}
    yield {"t": "progress", "done": idx + 1, "total": total}


# =====================================================================
#  Main generator
# =====================================================================
def iter_generate(kind, dfs, template_path, out_dir):
    """Generator populasi dokumen BAST (PPL atau PML).

    Yields event dicts:
      {"t": "log",      "level": str, "msg": str}
      {"t": "progress",  "done": int, "total": int}
      {"t": "file",      "path": str}
      {"t": "done",      "generated": [str], "skipped": [str]}
    """
    os.makedirs(out_dir, exist_ok=True)

    df_mitra = dfs.get(SHEET_NAME)
    if df_mitra is None or df_mitra.empty:
        yield {"t": "log", "level": "ERROR",
               "msg": f"Sheet '{SHEET_NAME}' kosong atau tidak ada."}
        yield {"t": "done", "generated": [], "skipped": []}
        return

    df_target = df_mitra[
        df_mitra["jabatan"].str.strip().str.lower() == kind
    ].copy()
    if df_target.empty:
        yield {"t": "log", "level": "WARN",
               "msg": f"Tidak ada data untuk {kind.upper()} "
                      f"(sheet '{SHEET_NAME}' "
                      f"kolom 'jabatan' tidak mengandung '{kind}')."}
        yield {"t": "done", "generated": [], "skipped": []}
        return

    df_tugas = dfs.get("alokasi_tugas", pd.DataFrame())
    df_kec = dfs.get("kode_kecamatan", pd.DataFrame())

    kec_map = {}
    if not df_kec.empty:
        for _, r in df_kec.iterrows():
            k = _norm(r.get("kode", ""))
            v = _norm(r.get("kecamatan", ""))
            if k:
                kec_map[k] = v

    total = len(df_target)
    kind_label = "PPL" if kind == "ppl" else "PML"
    generated, skipped = [], []

    yield {"t": "log", "level": "STEP",
           "msg": f"Memproses {total} {kind_label}..."}

    for idx, (_, row) in enumerate(df_target.iterrows()):
        nik = _norm(row.get("nik", ""))
        if not nik:
            yield {"t": "log", "level": "WARN",
                   "msg": f"   Baris {idx + 1}: NIK kosong - dilewati"}
            skipped.append(f"baris {idx + 1} (NIK kosong)")
            yield {"t": "progress", "done": idx + 1, "total": total}
            continue

        for ev in _generate_one_doc(
            row, kind, df_tugas, df_mitra, kec_map,
            template_path, out_dir, idx, total,
        ):
            if ev["t"] == "file":
                generated.append(ev["path"])
            yield ev

    yield {"t": "log", "level": "STEP", "msg": "=" * 50}
    yield {"t": "log",
           "level": "OK" if generated else "ERROR",
           "msg": f"Selesai: {len(generated)} dokumen berhasil, "
                  f"{len(skipped)} dilewati."}
    if skipped:
        yield {"t": "log", "level": "WARN",
               "msg": "Dilewati: " + ", ".join(skipped)}
    yield {"t": "done", "generated": generated, "skipped": skipped}
