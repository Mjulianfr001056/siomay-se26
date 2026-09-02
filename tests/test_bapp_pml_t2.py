"""Regression tests for the BAPP PML Termin 2 generator."""

import os
import tempfile
import unittest
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

from src import bapp_pml_t2


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "template" / "03. Template BAPP T2 PML.docx"


def _document_xml_text(path):
    """Return text-bearing XML from all Word document parts."""
    with zipfile.ZipFile(path) as archive:
        return "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


class BappPmlTermin2Tests(unittest.TestCase):
    def test_replace_handles_split_runs_and_surrounding_text(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Nomor: {{no_")
        paragraph.add_run("spk}}; bukti: {{bukti_dukung_")
        paragraph.add_run("bapp_t2}} selesai.")

        bapp_pml_t2.replace_text_preserving_runs(
            doc,
            {
                "{{no_spk}}": "SPK-PML-T2-001",
                "{{bukti_dukung_bapp_t2}}": "",
            },
        )

        self.assertEqual(paragraph.text, "Nomor: SPK-PML-T2-001; bukti:  selesai.")

    def test_actual_template_removes_every_placeholder_without_image_link(self):
        values = {
            "nik": "6304000000000002",
            "nama_lengkap": "PETUGAS UJI PML TERMIN DUA",
            "no_spk": "B-002/SPK-PML/2026",
            "no_urut_bapp_t2": "21",
            "jml_sls_t2": "31",
            "bukti_dukung_bapp_t2": "",
        }
        dfs = {bapp_pml_t2.SHEET_NAME: pd.DataFrame([values])}

        with tempfile.TemporaryDirectory() as output_dir:
            events = list(
                bapp_pml_t2.iter_generate(
                    dfs, os.fspath(TEMPLATE_PATH), output_dir
                )
            )
            output_path = next(
                event["path"] for event in events if event["t"] == "file"
            )
            xml_text = _document_xml_text(output_path)

        for value in values.values():
            if value:
                self.assertIn(value, xml_text)
        self.assertNotIn("{{bukti_dukung_bapp_t2}}", xml_text)
        self.assertNotIn("{{", xml_text)
        self.assertNotIn("}}", xml_text)


if __name__ == "__main__":
    unittest.main()