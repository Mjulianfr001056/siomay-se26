"""Regression tests for renamed SPP Termin I number input and placeholder."""
import os
import tempfile
import unittest
import zipfile

import pandas as pd
from docx import Document

from src import spp
from src.spp_t2 import _template_fields
from src.workflow import get_document_by_id


def _document_xml_text(path):
    with zipfile.ZipFile(path) as archive:
        return "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


class SppTermin1InputNameTests(unittest.TestCase):
    def test_catalog_labels_first_spp_as_termin_1(self):
        self.assertEqual(get_document_by_id("spp_ppl").label, "SPP PPL Termin 1")
        self.assertEqual(get_document_by_id("spp_ppl").group, "SPP Termin 1")
        self.assertEqual(get_document_by_id("spp_pml").label, "SPP PML Termin 1")
        self.assertEqual(get_document_by_id("spp_pml").group, "SPP Termin 1")

    def test_generator_uses_no_urut_spp_t1_for_its_matching_placeholder(self):
        with tempfile.TemporaryDirectory() as directory:
            template_path = os.path.join(directory, "template.docx")
            document = Document()
            document.add_paragraph("{{no_urut_spp_t1}}")
            document.add_paragraph("{{nama_lengkap}}")
            document.add_paragraph("{{nik}}")
            document.add_paragraph("{{no_spk}}")
            document.add_paragraph("{{jml_usaha}}")
            document.add_paragraph("{{jml_usaha_min}}")
            document.add_paragraph("{{persentase}}")
            document.save(template_path)

            dfs = {
                spp.SHEET_DATA_MITRA: pd.DataFrame([{
                    "nik": "6304000000000001",
                    "nama_lengkap": "PETUGAS UJI PPL",
                    "jabatan": "PPL",
                }]),
                spp.SHEET_NO_SPK: pd.DataFrame([{
                    "nik": "6304000000000001",
                    "no_spk": "SPK-001/2026",
                    "no_urut_spp_t1": "11",
                }]),
                spp.SHEET_ALOKASI: pd.DataFrame([{
                    "nik_ppl": "6304000000000001",
                    "nik_pml": "6304000000000002",
                    "target": "10",
                    "capaian": "10",
                    "persentase": "100",
                }]),
            }
            events = list(spp.iter_generate("ppl", dfs, template_path, directory))
            output_path = next(event["path"] for event in events if event["t"] == "file")
            xml_text = _document_xml_text(output_path)

        self.assertIn("11", xml_text)
        self.assertNotIn("{{no_urut_spp_t1}}", xml_text)

    def test_bundled_templates_use_no_urut_spp_t1(self):
        for document_id in ("spp_ppl", "spp_pml"):
            with self.subTest(document=document_id):
                fields = _template_fields(
                    get_document_by_id(document_id).builtin_template_path
                )
                self.assertIn("no_urut_spp_t1", fields)
                self.assertNotIn("no_input_spp_t1", fields)

    def test_custom_fields_replace_exact_text_blank_and_cannot_override_builtins(self):
        with tempfile.TemporaryDirectory() as directory:
            template_path = os.path.join(directory, "template.docx")
            document = Document()
            document.add_paragraph(
                "{{custom_code}}|{{blank_custom}}|{{nama_lengkap}}"
            )
            document.save(template_path)
            dfs = {
                spp.SHEET_DATA_MITRA: pd.DataFrame([{
                    "nik": "001",
                    "nama_lengkap": "NAMA RESMI",
                    "jabatan": "PPL",
                    "custom_code": "001",
                    "blank_custom": "",
                    # A custom/data value may not override the built-in mapping.
                    "{{nama_lengkap}}": "SALAH",
                }]),
                spp.SHEET_NO_SPK: pd.DataFrame([{
                    "nik": "001", "no_spk": "SPK", "no_urut_spp_t1": "1",
                }]),
                spp.SHEET_ALOKASI: pd.DataFrame([{
                    "nik_ppl": "001", "nik_pml": "002",
                    "target": "1", "capaian": "1", "persentase": "100",
                }]),
            }
            events = list(spp.iter_generate("ppl", dfs, template_path, directory))
            output_path = next(event["path"] for event in events if event["t"] == "file")
            output = Document(output_path)

        self.assertEqual(output.paragraphs[0].text, "001||NAMA RESMI")


if __name__ == "__main__":
    unittest.main()