"""Tests for BAPP Termin 2 evidence-image layout modes."""

import io
import inspect
import unittest
from unittest.mock import patch

from docx import Document
from PIL import Image

from src import bapp_pml_t2, bapp_ppl_t2
from utils import evidence


MODULES = (bapp_ppl_t2, bapp_pml_t2)


def _downloaded_image(_file_id):
    image = Image.new("RGB", (1200, 800), color=(40, 100, 180))
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    stream.seek(0)
    return stream, image


def _evidence_image(_file_id):
    stream, image = _downloaded_image(_file_id)
    size = image.size
    image.close()
    return [("image", stream, size)]


def _png_item(kind="image", color=(40, 100, 180)):
    image = Image.new("RGB", (1200, 800), color=color)
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    image.close()
    stream.seek(0)
    return kind, stream, (1200, 800)


def _sized_png_item(size, kind="image", color=(40, 100, 180)):
    image = Image.new("RGB", size, color=color)
    stream = io.BytesIO()
    image.save(stream, format="PNG")
    image.close()
    stream.seek(0)
    return kind, stream, size


def _document_with_placeholder(module):
    doc = Document()
    paragraph = doc.add_paragraph()
    paragraph.add_run("Sebelum ")
    paragraph.add_run(module.BUKTI_PLACEHOLDER[:14])
    paragraph.add_run(module.BUKTI_PLACEHOLDER[14:] + " sesudah")
    trailing_paragraph = doc.add_paragraph("KONTEN SETELAH PLACEHOLDER")
    return doc, paragraph, trailing_paragraph


class BappTermin2ImageLayoutTests(unittest.TestCase):
    def test_bapp_generator_orientation_defaults_to_portrait(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                parameters = inspect.signature(module.iter_generate).parameters
                self.assertEqual(
                    parameters["image_orientation"].default,
                    module.IMAGE_ORIENTATION_PORTRAIT,
                )

    def test_orientation_defaults_to_portrait_for_old_api_calls(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                with patch.object(
                    module, "_download_drive_evidence", side_effect=_evidence_image
                ):
                    module.insert_gdrive_images(
                        doc,
                        "https://drive.google.com/file/d/image-default/view",
                        image_layout=module.IMAGE_LAYOUT_DEDICATED_PAGES,
                    )
                shape = doc.inline_shapes[0]
                self.assertLessEqual(shape.width.inches,
                                     module.DEDICATED_MAX_WIDTH_IN)
                self.assertLessEqual(shape.height.inches,
                                     module.DEDICATED_MAX_HEIGHT_IN)

    def test_landscape_rotates_png_pixels_clockwise_without_mutating_source(self):
        image = Image.new("RGB", (2, 3), color=(0, 0, 0))
        image.putpixel((0, 0), (255, 0, 0))
        image.putpixel((1, 2), (0, 255, 0))
        source = io.BytesIO()
        image.save(source, format="PNG")
        original_bytes = source.getvalue()
        image.close()

        _, rotated_stream, rotated_size = evidence._prepare_dedicated_item(
            ("image", source, (2, 3)), evidence.IMAGE_ORIENTATION_LANDSCAPE
        )

        self.assertEqual(rotated_size, (3, 2))
        self.assertEqual(source.getvalue(), original_bytes)
        self.assertIsNot(rotated_stream, source)
        with Image.open(rotated_stream) as rotated:
            self.assertEqual(rotated.size, (3, 2))
            self.assertEqual(rotated.getpixel((2, 0)), (255, 0, 0))
            self.assertEqual(rotated.getpixel((0, 1)), (0, 255, 0))

    def test_landscape_always_rotates_wide_and_tall_images(self):
        for size in ((1200, 800), (800, 1200)):
            with self.subTest(size=size):
                item = _sized_png_item(size)
                _, result_stream, result_size = evidence._prepare_dedicated_item(
                    item, evidence.IMAGE_ORIENTATION_LANDSCAPE
                )
                self.assertIsNot(result_stream, item[1])
                self.assertEqual(result_size, (size[1], size[0]))

    def test_automatic_rotation_rules_cover_wide_tall_and_square_images(self):
        cases = [
            ((1200, 800), False, (1200, 800)),
            ((800, 1200), True, (1200, 800)),
            ((900, 900), False, (900, 900)),
        ]
        for size, rotates, expected_size in cases:
            with self.subTest(size=size):
                item = _sized_png_item(size)
                _, result_stream, result_size = evidence._prepare_dedicated_item(
                    item, evidence.IMAGE_ORIENTATION_AUTOMATIC
                )
                self.assertEqual(result_size, expected_size)
                self.assertEqual(result_stream is item[1], not rotates)

    def test_new_orientations_use_actual_page_area_and_first_unit_reserve(self):
        doc = Document()
        section = doc.sections[-1]
        content_width = (
            section.page_width - section.left_margin - section.right_margin
        ) / evidence.EMU_PER_INCH
        content_height = (
            section.page_height - section.top_margin - section.bottom_margin
        ) / evidence.EMU_PER_INCH

        later_box = evidence._dedicated_box(
            section, evidence.IMAGE_ORIENTATION_LANDSCAPE, False
        )
        first_box = evidence._dedicated_box(
            section, evidence.IMAGE_ORIENTATION_LANDSCAPE, True
        )
        square_width, square_height = evidence._fit_box(
            1000, 1000, *first_box
        )

        self.assertAlmostEqual(later_box[0], content_width)
        self.assertAlmostEqual(
            later_box[1], content_height - evidence.DEDICATED_TITLE_SPACE_IN
        )
        self.assertAlmostEqual(
            later_box[1] - first_box[1],
            evidence.DEDICATED_FIRST_UNIT_EXTRA_SPACE_IN,
        )
        self.assertAlmostEqual(square_width, square_height)
        self.assertAlmostEqual(square_width, min(first_box))

    def test_landscape_inserted_shape_stays_inside_actual_page_area(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                item = _sized_png_item((800, 1200))
                with patch.object(
                    module, "_download_drive_evidence", return_value=[item]
                ):
                    module.insert_gdrive_images(
                        doc,
                        "https://drive.google.com/file/d/image-landscape/view",
                        image_layout=module.IMAGE_LAYOUT_DEDICATED_PAGES,
                        image_orientation=module.IMAGE_ORIENTATION_LANDSCAPE,
                    )
                shape = doc.inline_shapes[0]
                box = evidence._dedicated_box(
                    doc.sections[-1], evidence.IMAGE_ORIENTATION_LANDSCAPE, True
                )
                self.assertLessEqual(shape.width.inches, box[0] + 0.001)
                self.assertLessEqual(shape.height.inches, box[1] + 0.001)

    def test_bapp_sequence_number_preserves_three_digit_form(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                self.assertEqual(module._format_no_urut_bapp_t2("001"), "001")
                # Excel commonly stores a displayed 001 as numeric 1.
                self.assertEqual(module._format_no_urut_bapp_t2(1), "001")
                self.assertEqual(module._format_no_urut_bapp_t2("21"), "021")
                self.assertEqual(module._format_no_urut_bapp_t2("AB-1"), "AB-1")

    def test_dedicated_pages_insert_every_image_without_grid(self):
        links = ",".join(
            f"https://drive.google.com/file/d/image-{number}/view"
            for number in range(1, 7)
        )

        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, placeholder_paragraph, trailing_paragraph = (
                    _document_with_placeholder(module)
                )
                with patch.object(
                    module, "_download_drive_evidence", side_effect=_evidence_image
                ) as downloader:
                    count, warnings = module.insert_gdrive_images(
                        doc,
                        links,
                        image_layout=module.IMAGE_LAYOUT_DEDICATED_PAGES,
                    )

                self.assertEqual(count, 6)
                self.assertEqual(warnings, [])
                self.assertEqual(downloader.call_count, 6)
                self.assertEqual(len(doc.inline_shapes), 6)
                self.assertEqual(len(doc.tables), 0)
                for shape in doc.inline_shapes:
                    self.assertLessEqual(
                        shape.width.inches, module.DEDICATED_MAX_WIDTH_IN
                    )
                    self.assertLessEqual(
                        shape.height.inches, module.DEDICATED_MAX_HEIGHT_IN
                    )
                self.assertNotIn(module.BUKTI_PLACEHOLDER, placeholder_paragraph.text)
                self.assertEqual(
                    len(doc.element.body.xpath('.//w:br[@w:type="page"]')), 5
                )
                titles = [p.text for p in doc.paragraphs]
                self.assertIn("BUKTI DUKUNG (1/6)", titles)
                self.assertIn("BUKTI DUKUNG (6/6)", titles)
                # Gambar pertama langsung mengikuti placeholder; seluruh bukti
                # tetap disisipkan sebelum konten template sesudah placeholder.
                body_elements = list(doc.element.body)
                first_title = next(
                    p for p in doc.paragraphs if p.text == "BUKTI DUKUNG (1/6)"
                )
                last_image_paragraph = next(
                    p for p in reversed(doc.paragraphs) if p._p.xpath(".//w:drawing")
                )
                self.assertEqual(
                    body_elements.index(first_title._p),
                    body_elements.index(placeholder_paragraph._p) + 1,
                )
                self.assertLess(
                    body_elements.index(last_image_paragraph._p),
                    body_elements.index(trailing_paragraph._p),
                )

    def test_grid_paginates_more_than_five_images(self):
        links = ",".join(
            f"https://drive.google.com/file/d/image-{number}/view"
            for number in range(1, 7)
        )

        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                with patch.object(
                    module, "_download_drive_evidence", side_effect=_evidence_image
                ) as downloader:
                    count, warnings = module.insert_gdrive_images(
                        doc, links, image_layout=module.IMAGE_LAYOUT_GRID
                    )

                self.assertEqual(count, 6)
                self.assertEqual(downloader.call_count, 6)
                self.assertEqual(len(doc.inline_shapes), 6)
                self.assertGreater(len(doc.tables), 0)
                self.assertEqual(warnings, [])
                self.assertEqual(
                    len(doc.element.body.xpath('.//w:br[@w:type="page"]')), 1
                )

    def test_pdf_pages_are_dedicated_then_images_resume_selected_grid(self):
        links = ",".join(
            f"https://drive.google.com/file/d/file-{number}/view"
            for number in range(1, 4)
        )

        def evidence(file_id):
            if file_id == "file-2":
                return [_png_item("pdf_page", (180, 40, 40)),
                        _png_item("pdf_page", (180, 80, 40))]
            return [_png_item()]

        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                with patch.object(
                    module, "_download_drive_evidence", side_effect=evidence
                ):
                    count, warnings = module.insert_gdrive_images(
                        doc, links, image_layout=module.IMAGE_LAYOUT_GRID
                    )

                self.assertEqual(count, 4)
                self.assertEqual(warnings, [])
                self.assertEqual(len(doc.inline_shapes), 4)
                self.assertEqual(len(doc.tables), 2)
                self.assertEqual(
                    len(doc.element.body.xpath('.//w:br[@w:type="page"]')), 3
                )
                pdf_titles = [p.text for p in doc.paragraphs
                              if p.text.startswith("BUKTI DUKUNG PDF")]
                self.assertEqual(pdf_titles,
                                 ["BUKTI DUKUNG PDF (2/4)",
                                  "BUKTI DUKUNG PDF (3/4)"])

    def test_invalid_layout_is_rejected(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                with self.assertRaisesRegex(ValueError, "tata letak gambar"):
                    module.insert_gdrive_images(
                        doc, "", image_layout="unknown-layout"
                    )

    def test_invalid_orientation_is_rejected_with_clear_error(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, _, _ = _document_with_placeholder(module)
                with self.assertRaisesRegex(ValueError, "Orientasi gambar"):
                    module.insert_gdrive_images(
                        doc, "", image_orientation="upside-down"
                    )

    def test_empty_links_still_remove_placeholder(self):
        for module in MODULES:
            with self.subTest(module=module.__name__):
                doc, paragraph, _ = _document_with_placeholder(module)
                count, warnings = module.insert_gdrive_images(doc, "")

                self.assertEqual(count, 0)
                self.assertEqual(warnings, [])
                self.assertEqual(paragraph.text, "Sebelum  sesudah")


if __name__ == "__main__":
    unittest.main()