"""Regression tests for SPP Termin 2 catalog and generator integration."""
import os
import tempfile
import unittest
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.dml.color import RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

from src import spp_t2
from src.workflow import DOCUMENT_TYPES, get_document_by_id


ROOT = Path(__file__).resolve().parents[1]


def _document_xml_text(path):
    with zipfile.ZipFile(path) as archive:
        return "".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in archive.namelist()
            if name.startswith("word/") and name.endswith(".xml")
        )


class SppTermin2WorkflowTests(unittest.TestCase):
    def test_catalog_places_spp_termin_ii_before_bast_with_shifted_assets(self):
        ids = [doc.id for doc in DOCUMENT_TYPES]
        self.assertLess(ids.index("spp_t2_ppl"), ids.index("bast_ppl"))
        self.assertLess(ids.index("spp_t2_pml"), ids.index("bast_pml"))
        self.assertEqual(get_document_by_id("spp_t2_ppl").label, "SPP PPL Termin 2")
        self.assertEqual(get_document_by_id("spp_t2_ppl").group, "SPP Termin 2")
        self.assertEqual(get_document_by_id("spp_t2_pml").label, "SPP PML Termin 2")
        self.assertEqual(get_document_by_id("spp_t2_pml").group, "SPP Termin 2")
        self.assertEqual(
            get_document_by_id("spp_t2_ppl").template_filename,
            "04. Template SPP T2 PPL.docx",
        )
        self.assertEqual(
            get_document_by_id("bast_ppl").template_filename,
            "05. Template BAST PPL.docx",
        )
        self.assertEqual(
            Path(get_document_by_id("bukti_terima").input_template_path).name,
            "06_input_bukti_terima_paket_internet.xlsx",
        )


class SppTermin2GeneratorTests(unittest.TestCase):
    @staticmethod
    def _dfs():
        return {
            spp_t2.SHEET_DATA_MITRA: pd.DataFrame([
                {
                    "nik": "6304000000000001",
                    "nama_lengkap": "PETUGAS UJI PPL",
                    "jabatan": "PPL",
                },
                {
                    "nik": "6304000000000002",
                    "nama_lengkap": "PETUGAS UJI PML",
                    "jabatan": "PML",
                },
            ]),
            spp_t2.SHEET_NO_SPK: pd.DataFrame([
                {
                    "nik": "6304000000000001",
                    "no_spk": "SPK-PPL-002/2026",
                    "no_urut_spp_t2": "21",
                },
                {
                    "nik": "6304000000000002",
                    "no_spk": "SPK-PML-002/2026",
                    "no_urut_spp_t2": "22",
                },
            ]),
            spp_t2.SHEET_ALOKASI: pd.DataFrame([{
                "nik_ppl": "6304000000000001",
                "nik_pml": "6304000000000002",
                "target": "10",
                "capaian": "8",
                "persentase": "80",
            }]),
        }

    def test_generator_replaces_template_fields_from_matching_input_columns(self):
        with tempfile.TemporaryDirectory() as directory:
            template_path = os.path.join(directory, "template.docx")
            document = Document()
            paragraph = document.add_paragraph()
            paragraph.add_run("{{nama_")
            paragraph.add_run("lengkap}}")
            document.add_paragraph("{{nik}}")
            document.add_paragraph("{{no_spk}}")
            document.add_paragraph("{{no_urut_spp_t2}}")
            document.save(template_path)

            events = list(spp_t2.iter_generate(
                "ppl", self._dfs(), template_path, directory
            ))
            output_path = next(event["path"] for event in events if event["t"] == "file")
            xml_text = _document_xml_text(output_path)

        self.assertIn("PETUGAS UJI PPL", xml_text)
        self.assertIn("6304000000000001", xml_text)
        self.assertIn("SPK-PPL-002/2026", xml_text)
        self.assertIn("21", xml_text)
        self.assertNotIn("{{", xml_text)

    def test_replaced_number_is_black_when_template_placeholder_is_red(self):
        with tempfile.TemporaryDirectory() as directory:
            template_path = os.path.join(directory, "template.docx")
            document = Document()
            paragraph = document.add_paragraph("Nomor: ")
            for text in ("{{no_", "urut", "_spp_t2}}"):
                run = paragraph.add_run(text)
                run.font.color.rgb = RGBColor(0xEE, 0x00, 0x00)
            document.save(template_path)

            events = list(spp_t2.iter_generate(
                "ppl", self._dfs(), template_path, directory
            ))
            output_path = next(
                event["path"] for event in events if event["t"] == "file"
            )
            output = Document(output_path)
            number_run = next(
                run for run in output.paragraphs[0].runs if "21" in run.text
            )

        self.assertEqual(number_run.font.color.rgb, RGBColor(0x00, 0x00, 0x00))

    def test_generated_pml_lampiran_table_is_centered(self):
        template_path = get_document_by_id("spp_t2_pml").builtin_template_path
        with tempfile.TemporaryDirectory() as directory:
            events = list(spp_t2.iter_generate(
                "pml", self._dfs(), template_path, directory
            ))
            output_path = next(
                event["path"] for event in events if event["t"] == "file"
            )
            output = Document(output_path)

        self.assertGreater(len(output.tables), 1)
        self.assertEqual(
            output.tables[1].alignment,
            WD_TABLE_ALIGNMENT.CENTER,
        )

    def test_validation_requires_columns_used_by_template(self):
        with tempfile.TemporaryDirectory() as directory:
            template_path = os.path.join(directory, "template.docx")
            document = Document()
            document.add_paragraph("{{nik}} {{nama_lengkap}} {{no_spk}} {{no_urut_spp_t2}}")
            document.save(template_path)
            input_path = os.path.join(directory, "input.xlsx")
            with pd.ExcelWriter(input_path) as writer:
                self._dfs()[spp_t2.SHEET_DATA_MITRA].to_excel(
                    writer, index=False, sheet_name=spp_t2.SHEET_DATA_MITRA
                )
                pd.DataFrame([{
                    "nik": "6304000000000001",
                    "no_spk": "SPK-002/2026",
                }]).to_excel(
                    writer, index=False, sheet_name=spp_t2.SHEET_NO_SPK
                )
                self._dfs()[spp_t2.SHEET_ALOKASI].to_excel(
                    writer, index=False, sheet_name=spp_t2.SHEET_ALOKASI
                )

            ok, errors, _ = spp_t2.validate_input(input_path, template_path)

        self.assertFalse(ok)
        self.assertIn("no_urut_spp_t2", errors[0])

    def test_bundled_templates_generate_with_no_unresolved_placeholders(self):
        for document_id in ("spp_t2_ppl", "spp_t2_pml"):
            document_type = get_document_by_id(document_id)
            ok, errors, _ = spp_t2.validate_input(
                document_type.input_template_path,
                document_type.builtin_template_path,
            )
            self.assertTrue(ok, errors)

            with self.subTest(document=document_id), tempfile.TemporaryDirectory() as directory:
                events = list(spp_t2.iter_generate(
                    document_type.kind,
                    self._dfs(),
                    document_type.builtin_template_path,
                    directory,
                ))
                output_paths = [
                    event["path"] for event in events if event["t"] == "file"
                ]

                self.assertTrue(output_paths)
                for output_path in output_paths:
                    xml_text = _document_xml_text(output_path)
                    self.assertNotIn("{{no_urut_spp_t2}}", xml_text)


if __name__ == "__main__":
    unittest.main()