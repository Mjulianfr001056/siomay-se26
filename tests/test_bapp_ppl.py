"""Regression tests for the BAPP PPL Termin 1 generator."""

import os
import tempfile
import unittest
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document

from src import bapp_pml, bapp_ppl


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_PATH = ROOT / "template" / "01. Template BAPP T1 PPL.docx"


def _document_xml_text(path):
    """Return all visible text nodes from every Word XML part."""
    with zipfile.ZipFile(path) as archive:
        chunks = []
        for name in archive.namelist():
            if name.startswith("word/") and name.endswith(".xml"):
                chunks.append(archive.read(name).decode("utf-8", errors="ignore"))
    return "".join(chunks)


class BappPplTermin1Tests(unittest.TestCase):
    def test_pml_replacement_visits_nested_tables_headers_and_footers(self):
        doc = Document()
        nested = doc.add_table(rows=1, cols=1).cell(0, 0).add_table(rows=1, cols=1)
        nested.cell(0, 0).text = "N={{custom}}"
        doc.sections[0].header.paragraphs[0].text = "H={{custom}}"
        doc.sections[0].footer.paragraphs[0].text = "F={{blank}}"

        bapp_pml.replace_text_preserving_runs(
            doc, {"{{custom}}": "001", "{{blank}}": ""}
        )

        self.assertEqual(nested.cell(0, 0).text, "N=001")
        self.assertEqual(doc.sections[0].header.paragraphs[0].text, "H=001")
        self.assertEqual(doc.sections[0].footer.paragraphs[0].text, "F=")

    def test_replace_handles_split_runs_and_surrounding_text(self):
        doc = Document()
        paragraph = doc.add_paragraph()
        paragraph.add_run("Nomor: {{no_")
        paragraph.add_run("spk}} selesai; {{nik}}.")

        bapp_ppl.replace_text_preserving_runs(
            doc,
            {"{{no_spk}}": "SPK-001", "{{nik}}": "6304000000000001"},
        )

        self.assertEqual(
            paragraph.text,
            "Nomor: SPK-001 selesai; 6304000000000001.",
        )

    def test_actual_template_populates_every_placeholder_without_image_link(self):
        values = {
            "nik": "6304000000000001",
            "nama_lengkap": "PETUGAS UJI PPL",
            "no_spk": "B-001/SPK-PPL/2026",
            "no_urut_bapp_t1": "17",
            "jml_sls_t1": "23",
            "bukti_dukung_bapp_t1": "",
        }
        dfs = {bapp_ppl.SHEET_NAME: pd.DataFrame([values])}

        with tempfile.TemporaryDirectory() as output_dir:
            events = list(
                bapp_ppl.iter_generate(dfs, os.fspath(TEMPLATE_PATH), output_dir)
            )
            output_path = next(event["path"] for event in events if event["t"] == "file")
            xml_text = _document_xml_text(output_path)

        for value in values.values():
            if value:
                self.assertIn(value, xml_text)
        self.assertNotIn("{{", xml_text)
        self.assertNotIn("}}", xml_text)


if __name__ == "__main__":
    unittest.main()