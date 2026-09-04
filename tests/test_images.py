"""Tests for shared JPEG/PNG/HEIC normalization."""

import io
import unittest

from docx import Document
from PIL import Image

from utils.images import (
    HAS_HEIF,
    HAS_PDF_RENDERER,
    image_bytes_to_png,
    pdf_bytes_to_png_pages,
)


class ImageNormalizationTests(unittest.TestCase):
    def test_jpeg_is_normalized_to_insertable_png(self):
        source = io.BytesIO()
        Image.new("RGB", (12, 8), "blue").save(source, format="JPEG")

        png_file, image = image_bytes_to_png(source.getvalue(), "image/jpeg")

        self.assertEqual(png_file.read(8), b"\x89PNG\r\n\x1a\n")
        self.assertEqual(image.size, (12, 8))
        png_file.seek(0)
        Document().add_paragraph().add_run().add_picture(png_file)

    @unittest.skipUnless(HAS_HEIF, "pillow-heif/native HEIF codec unavailable")
    def test_heic_is_decoded_to_insertable_png(self):
        source = io.BytesIO()
        Image.new("RGB", (11, 7), "green").save(source, format="HEIF")

        png_file, image = image_bytes_to_png(source.getvalue(), "image/heic")

        self.assertEqual(png_file.read(8), b"\x89PNG\r\n\x1a\n")
        self.assertEqual(image.size, (11, 7))
        png_file.seek(0)
        Document().add_paragraph().add_run().add_picture(png_file)

    def test_html_download_is_rejected_with_actionable_message(self):
        with self.assertRaisesRegex(RuntimeError, "halaman HTML"):
            image_bytes_to_png(b"<!doctype html><html></html>", "text/html")

    @unittest.skipUnless(HAS_PDF_RENDERER, "PyMuPDF unavailable")
    def test_multipage_pdf_is_rendered_to_insertable_png_pages(self):
        import pymupdf as fitz

        pdf = fitz.open()
        for text in ("Halaman satu", "Halaman dua"):
            page = pdf.new_page(width=300, height=500)
            page.insert_text((40, 60), text)
        raw_pdf = pdf.tobytes()
        pdf.close()

        pages = pdf_bytes_to_png_pages(raw_pdf)

        self.assertEqual(len(pages), 2)
        for stream, size in pages:
            self.assertEqual(stream.read(8), b"\x89PNG\r\n\x1a\n")
            self.assertGreater(size[0], 0)
            self.assertGreater(size[1], 0)
            stream.seek(0)
            Document().add_paragraph().add_run().add_picture(stream)


if __name__ == "__main__":
    unittest.main()