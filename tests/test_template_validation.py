"""Tests for validation of edited DOCX templates in Step 2."""

import os
import io
import tempfile
import unittest
from unittest.mock import Mock

from docx import Document
from PIL import Image

from src.document_generator import (
    extend_input_template,
    extract_template_placeholders,
    fill_row,
    insert_custom_url_images,
    row_placeholder_replacements,
    validate_custom_columns,
    validate_template_placeholders,
)
import openpyxl
import pandas as pd


class TemplatePlaceholderValidationTests(unittest.TestCase):
    def _save_document(self, directory, name, configure_document):
        document = Document()
        configure_document(document)
        path = os.path.join(directory, name)
        document.save(path)
        return path

    def test_extracts_placeholders_from_split_runs_tables_and_header(self):
        with tempfile.TemporaryDirectory() as directory:
            def configure(document):
                paragraph = document.add_paragraph()
                paragraph.add_run("{{nama_")
                paragraph.add_run("lengkap}}")
                document.add_table(rows=1, cols=1).cell(0, 0).text = "{{nik}}"
                document.sections[0].header.paragraphs[0].text = "{{no_spk}}"

            path = self._save_document(directory, "template.docx", configure)

            self.assertEqual(
                extract_template_placeholders(path),
                {"nama_lengkap", "nik", "no_spk"},
            )

    def test_accepts_template_with_exact_expected_placeholders(self):
        with tempfile.TemporaryDirectory() as directory:
            path = self._save_document(
                directory,
                "edited.docx",
                lambda document: document.add_paragraph(
                    "{{nik}} untuk {{nama_lengkap}}"
                ),
            )

            result = validate_template_placeholders(
                path, {"nama_lengkap", "nik"}
            )

            self.assertTrue(result["is_valid"])
            self.assertEqual(result["missing"], [])
            self.assertEqual(result["unexpected"], [])

    def test_rejects_missing_but_accepts_custom_placeholders(self):
        with tempfile.TemporaryDirectory() as directory:
            path = self._save_document(
                directory,
                "invalid.docx",
                lambda document: document.add_paragraph("{{nik}} {{jabatan}}"),
            )

            result = validate_template_placeholders(
                path, {"nama_lengkap", "nik"}
            )

            self.assertFalse(result["is_valid"])
            self.assertEqual(result["missing"], ["nama_lengkap"])
            self.assertEqual(result["unexpected"], ["jabatan"])
            self.assertEqual(result["custom"], ["jabatan"])

    def test_accepts_custom_placeholder_when_builtins_are_retained(self):
        with tempfile.TemporaryDirectory() as directory:
            path = self._save_document(
                directory, "custom.docx",
                lambda document: document.add_paragraph(
                    "{{nik}} {{nama_lengkap}} {{kode_custom}}"
                ),
            )
            result = validate_template_placeholders(
                path, {"nama_lengkap", "nik"}
            )
            self.assertTrue(result["is_valid"])
            self.assertEqual(result["custom"], ["kode_custom"])

    def test_ignores_malformed_placeholder_markers(self):
        with tempfile.TemporaryDirectory() as directory:
            path = self._save_document(
                directory, "malformed.docx",
                lambda document: document.add_paragraph(
                    "{{valid_1}} {{ invalid }} {{not-valid}} {{é}}"
                ),
            )
            self.assertEqual(extract_template_placeholders(path), {"valid_1"})

    def test_extracts_nested_table_header_table_and_footer_table(self):
        with tempfile.TemporaryDirectory() as directory:
            def configure(document):
                outer = document.add_table(rows=1, cols=1)
                outer.cell(0, 0).add_table(rows=1, cols=1).cell(0, 0).text = "{{nested}}"
                header = document.sections[0].header
                header.add_table(rows=1, cols=1, width=1).cell(0, 0).text = "{{header_table}}"
                footer = document.sections[0].footer
                footer.add_table(rows=1, cols=1, width=1).cell(0, 0).text = "{{footer_table}}"

            path = self._save_document(directory, "locations.docx", configure)
            self.assertEqual(
                extract_template_placeholders(path),
                {"nested", "header_table", "footer_table"},
            )

    def test_extracts_and_fills_all_header_and_footer_variants(self):
        with tempfile.TemporaryDirectory() as directory:
            def configure(document):
                section = document.sections[0]
                section.different_first_page_header_footer = True
                document.settings.odd_and_even_pages_header_footer = True
                stories = {
                    "header": section.header,
                    "footer": section.footer,
                    "first_header": section.first_page_header,
                    "first_footer": section.first_page_footer,
                    "even_header": section.even_page_header,
                    "even_footer": section.even_page_footer,
                }
                for field, story in stories.items():
                    story.paragraphs[0].text = "{{" + field + "}}"
                section.even_page_footer.add_table(
                    rows=1, cols=1, width=1,
                ).cell(0, 0).text = "{{even_footer_table}}"

            source = self._save_document(directory, "all-stories.docx", configure)
            fields = {
                "header", "footer", "first_header", "first_footer",
                "even_header", "even_footer", "even_footer_table",
            }
            self.assertEqual(extract_template_placeholders(source), fields)

            output = os.path.join(directory, "filled.docx")
            result = fill_row(source, {field: field.upper() for field in fields}, output)
            self.assertEqual(result["filled"], sorted(fields))
            self.assertEqual(result["unresolved"], [])
            self.assertEqual(extract_template_placeholders(output), set())

    def test_custom_column_matching_is_trimmed_but_case_sensitive(self):
        dfs = {"input": pd.DataFrame(columns=[" custom_code ", "ExactCase"])}
        self.assertEqual(
            validate_custom_columns(dfs, "input", ["custom_code", "ExactCase"]), []
        )
        errors = validate_custom_columns(dfs, "input", ["exactcase"])
        self.assertIn("exactcase", errors[0])

    def test_row_replacements_preserve_text_and_include_blank(self):
        row = pd.Series({"code": "001", "blank": "", "bad name": "ignored"})
        self.assertEqual(
            row_placeholder_replacements(row),
            {"{{code}}": "001", "{{blank}}": ""},
        )

    def test_custom_image_urls_use_evidence_layout_from_all_docx_stories(self):
        document = Document()
        paragraph = document.add_paragraph("Before ")
        paragraph.add_run("{{photo_")
        paragraph.add_run("custom}} after {{photo_custom}}")
        nested = document.add_table(rows=1, cols=1).cell(0, 0)
        nested.add_table(rows=1, cols=1).cell(0, 0).text = "{{photo_custom}}"
        document.sections[0].header.paragraphs[0].text = "{{photo_custom}}"

        image = Image.new("RGB", (120, 80), color="navy")
        stream = io.BytesIO()
        image.save(stream, format="PNG")
        image.close()
        stream.seek(0)
        downloaded = Image.open(stream)
        downloader = Mock(return_value=(stream, downloaded))

        consumed = insert_custom_url_images(
            document,
            pd.Series({"photo_custom": "https://example.test/photo"}),
            downloader=downloader,
        )

        self.assertEqual(consumed, {"{{photo_custom}}"})
        self.assertEqual(downloader.call_count, 1)
        drawing_count = len(document.element.xpath(".//w:drawing"))
        drawing_count += len(
            document.sections[0].header._element.xpath(".//w:drawing")
        )
        self.assertEqual(drawing_count, 4)
        self.assertNotIn("{{photo_custom}}", paragraph.text)
        self.assertEqual(paragraph.text, "Before  after ")

    def test_custom_image_uses_selected_dedicated_orientation(self):
        document = Document()
        placeholder = document.add_paragraph("{{photo_custom}}")
        image = Image.new("RGB", (80, 120), color="navy")
        stream = io.BytesIO()
        image.save(stream, format="PNG")
        image.close()
        stream.seek(0)

        consumed = insert_custom_url_images(
            document,
            pd.Series({"photo_custom": "https://example.test/photo"}),
            image_layout="dedicated_pages",
            image_orientation="landscape",
            downloader=Mock(return_value=(stream, Image.open(stream))),
        )

        self.assertEqual(consumed, {"{{photo_custom}}"})
        self.assertEqual(placeholder.text, "")
        self.assertEqual(len(document.inline_shapes), 1)
        shape = document.inline_shapes[0]
        self.assertGreater(shape.width, shape.height)
        self.assertFalse(any(
            p.text.startswith("BUKTI DUKUNG") for p in document.paragraphs
        ))

    def test_custom_non_image_or_failed_url_is_left_for_text_fallback(self):
        for side_effect in (RuntimeError("not an image"), ValueError("download failed")):
            with self.subTest(error=type(side_effect).__name__):
                document = Document()
                document.add_paragraph("Source: {{reference_custom}}")
                url = "https://example.test/document.pdf"
                replacements = row_placeholder_replacements(
                    pd.Series({"reference_custom": url})
                )

                consumed = insert_custom_url_images(
                    document,
                    pd.Series({"reference_custom": url}),
                    downloader=Mock(side_effect=side_effect),
                )
                from src.spp import replace_text_preserving_runs
                replace_text_preserving_runs(document, replacements)

                self.assertEqual(consumed, set())
                self.assertEqual(document.paragraphs[0].text, "Source: " + url)
                self.assertEqual(len(document.inline_shapes), 0)

    def test_custom_plain_text_does_not_trigger_download(self):
        document = Document()
        document.add_paragraph("{{note_custom}}")
        downloader = Mock()
        row = pd.Series({"note_custom": "ordinary text"})

        consumed = insert_custom_url_images(document, row, downloader=downloader)

        self.assertEqual(consumed, set())
        downloader.assert_not_called()

    def test_extends_copy_with_text_format_without_modifying_source(self):
        with tempfile.TemporaryDirectory() as directory:
            source = os.path.join(directory, "source.xlsx")
            output = os.path.join(directory, "output.xlsx")
            workbook = openpyxl.Workbook()
            worksheet = workbook.active
            worksheet.title = "input"
            worksheet.append(["nik"])
            worksheet.append(["001"])
            workbook.save(source)
            workbook.close()

            extend_input_template(source, output, "input", ["custom_code"])

            original = openpyxl.load_workbook(source)
            generated = openpyxl.load_workbook(output)
            self.assertEqual(original["input"].max_column, 1)
            self.assertEqual(generated["input"]["B1"].value, "custom_code")
            self.assertEqual(generated["input"]["B1"].number_format, "@")
            self.assertEqual(generated["input"]["B2"].number_format, "@")
            original.close()
            generated.close()


if __name__ == "__main__":
    unittest.main()