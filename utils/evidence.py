"""Shared DOCX layout engine for image and PDF evidence from Google Drive."""

import io

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image

IMAGE_LAYOUT_GRID = "grid"
IMAGE_LAYOUT_DEDICATED_PAGES = "dedicated_pages"
IMAGE_LAYOUTS = {IMAGE_LAYOUT_GRID, IMAGE_LAYOUT_DEDICATED_PAGES}
IMAGE_ORIENTATION_PORTRAIT = "portrait"
IMAGE_ORIENTATION_LANDSCAPE = "landscape"
IMAGE_ORIENTATION_AUTOMATIC = "automatic"
IMAGE_ORIENTATIONS = {
    IMAGE_ORIENTATION_PORTRAIT,
    IMAGE_ORIENTATION_LANDSCAPE,
    IMAGE_ORIENTATION_AUTOMATIC,
}
GRID_LAYOUTS = {1: [1], 2: [2], 3: [2, 1], 4: [2, 2], 5: [3, 2]}
GRID_MAX_WIDTH_IN = 9.2
GRID_MAX_HEIGHT_IN = 4.6
GRID_GAP_IN = 0.12
DEDICATED_MAX_WIDTH_IN = 7.5
DEDICATED_MAX_HEIGHT_IN = 4.0
DEDICATED_TITLE_SPACE_IN = 0.75
DEDICATED_FIRST_UNIT_EXTRA_SPACE_IN = 0.25
EMU_PER_INCH = 914400


def _fit_box(img_w, img_h, box_w, box_h):
    aspect = img_w / img_h
    target_w, target_h = box_w, box_w / aspect
    if target_h > box_h:
        target_h, target_w = box_h, box_h * aspect
    return target_w, target_h


def _page_break(doc, anchor):
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)
    anchor.addnext(paragraph._p)
    return paragraph._p


def _rotate_clockwise_png(stream):
    """Return a new PNG stream and size after a pixel-level clockwise rotation."""
    stream.seek(0)
    with Image.open(stream) as image:
        rotated = image.rotate(-90, expand=True)
        output = io.BytesIO()
        rotated.save(output, format="PNG")
        size = rotated.size
        rotated.close()
    output.seek(0)
    return output, size


def _prepare_dedicated_item(item, image_orientation):
    kind, stream, size = item
    rotate = (
        image_orientation == IMAGE_ORIENTATION_LANDSCAPE
        or (
            image_orientation == IMAGE_ORIENTATION_AUTOMATIC
            and size[1] > size[0]
        )
    )
    if not rotate:
        return kind, stream, size
    rotated_stream, rotated_size = _rotate_clockwise_png(stream)
    return kind, rotated_stream, rotated_size


def _dedicated_box(section, image_orientation, is_first_unit):
    content_w = (
        section.page_width - section.left_margin - section.right_margin
    ) / EMU_PER_INCH
    content_h = (
        section.page_height - section.top_margin - section.bottom_margin
    ) / EMU_PER_INCH
    if image_orientation == IMAGE_ORIENTATION_PORTRAIT:
        return (
            max(min(content_w, DEDICATED_MAX_WIDTH_IN), 1.0),
            max(
                min(content_h - DEDICATED_TITLE_SPACE_IN,
                    DEDICATED_MAX_HEIGHT_IN),
                1.0,
            ),
        )

    first_unit_space = (
        DEDICATED_FIRST_UNIT_EXTRA_SPACE_IN if is_first_unit else 0.0
    )
    return (
        max(content_w, 1.0),
        max(content_h - DEDICATED_TITLE_SPACE_IN - first_unit_space, 1.0),
    )


def _insert_dedicated(doc, anchor, item, number, total, image_orientation,
                      is_first_unit=False, show_title=True):
    kind, stream, size = _prepare_dedicated_item(item, image_orientation)
    section = doc.sections[-1]
    box_w, box_h = _dedicated_box(
        section, image_orientation, is_first_unit
    )

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    width, height = _fit_box(*size, box_w, box_h)
    stream.seek(0)
    paragraph.add_run().add_picture(
        stream, width=Inches(width), height=Inches(height)
    )
    if show_title:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_after = Pt(8)
        label = "BUKTI DUKUNG PDF" if kind == "pdf_page" else "BUKTI DUKUNG"
        run = title.add_run(f"{label} ({number}/{total})")
        run.bold = True
        run.font.size = Pt(12)
        anchor.addnext(title._p)
        title._p.addnext(paragraph._p)
    else:
        anchor.addnext(paragraph._p)
    return paragraph._p


def _remove_table_borders(table):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "none")
        element.set(qn("w:sz"), "0")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "auto")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def _set_cell_width(cell, width):
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(width * 1440)))


def _insert_grid(doc, anchor, images):
    layout = GRID_LAYOUTS[len(images)]
    row_h = (GRID_MAX_HEIGHT_IN - GRID_GAP_IN * (len(layout) - 1)) / len(layout)
    image_index = 0
    for columns in layout:
        col_w = (GRID_MAX_WIDTH_IN - GRID_GAP_IN * (columns - 1)) / columns
        table = doc.add_table(rows=1, cols=columns)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        _remove_table_borders(table)
        for cell in table.rows[0].cells:
            _set_cell_width(cell, col_w)
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            if image_index < len(images):
                _, stream, size = images[image_index]
                width, height = _fit_box(*size, col_w, row_h)
                stream.seek(0)
                paragraph.add_run().add_picture(
                    stream, width=Inches(width), height=Inches(height)
                )
                image_index += 1
        anchor.addnext(table._tbl)
        anchor = table._tbl
    return anchor


def insert_evidence_items(doc, target, items, image_layout,
                          image_orientation=IMAGE_ORIENTATION_PORTRAIT,
                          show_titles=True):
    """Lay out already-downloaded evidence immediately after *target*.

    Ordinary images follow the selected grid/dedicated layout. Rendered PDF
    pages are always placed on dedicated pages, matching built-in evidence.
    """
    if image_layout not in IMAGE_LAYOUTS:
        raise ValueError(f"Mode tata letak gambar tidak dikenal: {image_layout}")
    if image_orientation not in IMAGE_ORIENTATIONS:
        raise ValueError(
            "Orientasi gambar tidak dikenal: " + str(image_orientation)
        )
    if not items:
        return 0

    # Build page units while preserving source order. PDF pages always form a
    # dedicated unit; ordinary images use the selected mode and resume after it.
    units = []
    pending_images = []

    def flush_images():
        while pending_images:
            size = 5 if image_layout == IMAGE_LAYOUT_GRID else 1
            units.append(("grid" if size == 5 else "dedicated",
                          pending_images[:size]))
            del pending_images[:size]

    for item in items:
        if item[0] == "pdf_page":
            flush_images()
            units.append(("dedicated", [item]))
        else:
            pending_images.append(item)
    flush_images()

    anchor = target._p
    evidence_number = 0
    total = len(items)
    for unit_index, (unit_kind, unit_items) in enumerate(units):
        if unit_index:
            anchor = _page_break(doc, anchor)
        if unit_kind == "grid":
            anchor = _insert_grid(doc, anchor, unit_items)
            evidence_number += len(unit_items)
        else:
            evidence_number += 1
            anchor = _insert_dedicated(
                doc, anchor, unit_items[0], evidence_number, total,
                image_orientation, is_first_unit=(unit_index == 0),
                show_title=show_titles,
            )
    return len(items)


def insert_evidence(doc, links_str, placeholder, image_layout, extract_file_id,
                    replace_text, evidence_downloader,
                    image_orientation=IMAGE_ORIENTATION_PORTRAIT):
    """Insert ordered images/PDF pages; PDF pages are always dedicated pages."""
    if image_layout not in IMAGE_LAYOUTS:
        raise ValueError(f"Mode tata letak gambar tidak dikenal: {image_layout}")
    if image_orientation not in IMAGE_ORIENTATIONS:
        raise ValueError(
            "Orientasi gambar tidak dikenal: " + str(image_orientation)
        )
    target = next((p for p in doc.paragraphs if placeholder in p.text), None)
    if target is None:
        return 0, ["Template tidak memiliki placeholder " + placeholder]
    replace_text(doc, {placeholder: ""})
    if not links_str or not str(links_str).strip():
        return 0, []

    warnings = []
    items = []
    for link in (value.strip() for value in str(links_str).split(",")):
        if not link:
            continue
        file_id = extract_file_id(link)
        if not file_id:
            warnings.append("Tautan tidak dikenali: " + link)
            continue
        try:
            items.extend(evidence_downloader(file_id))
        except Exception as exc:
            message = str(exc)
            if "403" in message or "forbidden" in message.lower():
                warnings.append(f"Akses ditolak (403) untuk {file_id}")
            elif "404" in message:
                warnings.append(f"File {file_id} tidak ditemukan")
            else:
                warnings.append(f"Gagal memuat {file_id}: {message}")

    if not items:
        return 0, warnings
    return insert_evidence_items(
        doc, target, items, image_layout, image_orientation
    ), warnings